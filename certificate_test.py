#!/usr/bin/env python3
"""
Certificate Generation Flow Test Suite
Tests the complete end-to-end certificate generation flow as requested:

Phase 1: Setup (Admin)
- Login as admin
- Create a program
- Create a company
- Create a session with the program and company
- Create a participant user
- Add participant to the session

Phase 2: Participant Access
- Login as participant
- GET /sessions - Verify participant_access record is auto-created
- Check participant_access has default values

Phase 3: Feedback Submission
- As coordinator/admin, release feedback: POST /sessions/{session_id}/release-feedback
- As participant, submit feedback: POST /feedback/submit
- Verify participant_access.feedback_submitted = True

Phase 4: Certificate Generation (THE CRITICAL PART)
- POST /certificates/generate/{session_id}/{participant_id}
- Check response includes certificate_url and download_url
- GET the download_url - verify .docx file downloads
- Check file size > 0
- Verify file is a valid Word document
"""

import requests
import json
import sys
import os
import tempfile
from datetime import datetime
from pathlib import Path

# Configuration
BASE_URL = "https://vehiclelearn.preview.emergentagent.com/api"
ADMIN_EMAIL = "admin@example.com"
ADMIN_PASSWORD = "admin123"

class CertificateTestRunner:
    def __init__(self):
        self.admin_token = None
        self.participant_token = None
        self.program_id = None
        self.company_id = None
        self.session_id = None
        self.participant_id = None
        self.session = requests.Session()
        self.session.headers.update({'Content-Type': 'application/json'})
        
    def log(self, message, level="INFO"):
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        print(f"[{timestamp}] [{level}] {message}")
        
    def login_admin(self):
        """Phase 1: Login as admin and get authentication token"""
        self.log("Phase 1: Attempting admin login...")
        
        login_data = {
            "email": ADMIN_EMAIL,
            "password": ADMIN_PASSWORD
        }
        
        try:
            response = self.session.post(f"{BASE_URL}/auth/login", json=login_data)
            
            if response.status_code == 200:
                data = response.json()
                self.admin_token = data['access_token']
                self.log(f"✅ Admin login successful. User: {data['user']['full_name']} ({data['user']['role']})")
                return True
            else:
                self.log(f"❌ Admin login failed: {response.status_code} - {response.text}", "ERROR")
                return False
                
        except Exception as e:
            self.log(f"❌ Admin login error: {str(e)}", "ERROR")
            return False
    
    def create_program(self):
        """Phase 1: Create a program"""
        self.log("Phase 1: Creating program...")
        
        if not self.admin_token:
            self.log("❌ No admin token available", "ERROR")
            return False
            
        headers = {'Authorization': f'Bearer {self.admin_token}'}
        
        program_data = {
            "name": "Certificate Test Program",
            "description": "Program created for certificate generation testing",
            "pass_percentage": 70.0
        }
        
        try:
            response = self.session.post(f"{BASE_URL}/programs", json=program_data, headers=headers)
            
            if response.status_code == 200:
                data = response.json()
                self.program_id = data['id']
                self.log(f"✅ Program created successfully. ID: {self.program_id}")
                return True
            else:
                self.log(f"❌ Program creation failed: {response.status_code} - {response.text}", "ERROR")
                return False
                
        except Exception as e:
            self.log(f"❌ Program creation error: {str(e)}", "ERROR")
            return False
    
    def create_company(self):
        """Phase 1: Create a company"""
        self.log("Phase 1: Creating company...")
        
        if not self.admin_token:
            self.log("❌ No admin token available", "ERROR")
            return False
            
        headers = {'Authorization': f'Bearer {self.admin_token}'}
        
        company_data = {
            "name": "Certificate Test Company Ltd"
        }
        
        try:
            response = self.session.post(f"{BASE_URL}/companies", json=company_data, headers=headers)
            
            if response.status_code == 200:
                data = response.json()
                self.company_id = data['id']
                self.log(f"✅ Company created successfully. ID: {self.company_id}")
                return True
            else:
                self.log(f"❌ Company creation failed: {response.status_code} - {response.text}", "ERROR")
                return False
                
        except Exception as e:
            self.log(f"❌ Company creation error: {str(e)}", "ERROR")
            return False
    
    def create_participant_user(self):
        """Phase 1: Create a participant user"""
        self.log("Phase 1: Creating participant user...")
        
        if not self.admin_token:
            self.log("❌ No admin token available", "ERROR")
            return False
            
        headers = {'Authorization': f'Bearer {self.admin_token}'}
        
        participant_data = {
            "email": "certificate.participant@example.com",
            "password": "participant123",
            "full_name": "Certificate Test Participant",
            "id_number": "CERT001",
            "role": "participant",
            "location": "Certificate Test Location"
        }
        
        try:
            response = self.session.post(f"{BASE_URL}/auth/register", json=participant_data, headers=headers)
            
            if response.status_code == 200:
                data = response.json()
                self.participant_id = data['id']
                self.log(f"✅ Participant user created successfully. ID: {self.participant_id}")
                return True
            elif response.status_code == 400 and "User already exists" in response.text:
                # Get existing user ID by logging in
                login_data = {
                    "email": "certificate.participant@example.com",
                    "password": "participant123"
                }
                response = self.session.post(f"{BASE_URL}/auth/login", json=login_data)
                if response.status_code == 200:
                    self.participant_id = response.json()['user']['id']
                    self.log(f"✅ Using existing participant user. ID: {self.participant_id}")
                    return True
                else:
                    self.log("❌ Failed to get existing participant ID", "ERROR")
                    return False
            else:
                self.log(f"❌ Participant creation failed: {response.status_code} - {response.text}", "ERROR")
                return False
                
        except Exception as e:
            self.log(f"❌ Participant creation error: {str(e)}", "ERROR")
            return False
    
    def create_session(self):
        """Phase 1: Create a session with the program and company"""
        self.log("Phase 1: Creating session with program and company...")
        
        if not self.admin_token or not self.program_id or not self.company_id or not self.participant_id:
            self.log("❌ Missing required data (admin token, program ID, company ID, or participant ID)", "ERROR")
            return False
            
        headers = {'Authorization': f'Bearer {self.admin_token}'}
        
        session_data = {
            "name": "Certificate Test Session",
            "program_id": self.program_id,
            "company_id": self.company_id,
            "location": "Certificate Test Location",
            "start_date": "2024-01-01",
            "end_date": "2024-01-31",
            "participant_ids": [self.participant_id]
        }
        
        try:
            response = self.session.post(f"{BASE_URL}/sessions", json=session_data, headers=headers)
            
            if response.status_code == 200:
                data = response.json()
                self.session_id = data['id']
                self.log(f"✅ Session created successfully. ID: {self.session_id}")
                self.log(f"   Program ID: {data['program_id']}")
                self.log(f"   Company ID: {data['company_id']}")
                self.log(f"   Participants: {len(data['participant_ids'])}")
                return True
            else:
                self.log(f"❌ Session creation failed: {response.status_code} - {response.text}", "ERROR")
                return False
                
        except Exception as e:
            self.log(f"❌ Session creation error: {str(e)}", "ERROR")
            return False
    
    def login_participant(self):
        """Phase 2: Login as participant"""
        self.log("Phase 2: Logging in as participant...")
        
        login_data = {
            "email": "certificate.participant@example.com",
            "password": "participant123"
        }
        
        try:
            response = self.session.post(f"{BASE_URL}/auth/login", json=login_data)
            
            if response.status_code == 200:
                data = response.json()
                self.participant_token = data['access_token']
                self.log(f"✅ Participant login successful. User: {data['user']['full_name']} ({data['user']['role']})")
                return True
            else:
                self.log(f"❌ Participant login failed: {response.status_code} - {response.text}", "ERROR")
                return False
                
        except Exception as e:
            self.log(f"❌ Participant login error: {str(e)}", "ERROR")
            return False
    
    def verify_participant_access_auto_created(self):
        """Phase 2: Verify participant_access record is auto-created when getting sessions"""
        self.log("Phase 2: Verifying participant_access auto-creation...")
        
        if not self.participant_token:
            self.log("❌ No participant token available", "ERROR")
            return False
            
        headers = {'Authorization': f'Bearer {self.participant_token}'}
        
        try:
            # GET /sessions should auto-create participant_access records
            response = self.session.get(f"{BASE_URL}/sessions", headers=headers)
            
            if response.status_code == 200:
                sessions = response.json()
                self.log(f"✅ Sessions retrieved successfully. Count: {len(sessions)}")
                
                # Find our test session
                test_session = None
                for session in sessions:
                    if session['id'] == self.session_id:
                        test_session = session
                        break
                
                if test_session:
                    self.log("✅ Test session found in participant's sessions")
                    return True
                else:
                    self.log("❌ Test session not found in participant's sessions", "ERROR")
                    return False
                    
            else:
                self.log(f"❌ Get sessions failed: {response.status_code} - {response.text}", "ERROR")
                return False
                
        except Exception as e:
            self.log(f"❌ Get sessions error: {str(e)}", "ERROR")
            return False
    
    def check_participant_access_defaults(self):
        """Phase 2: Check participant_access has default values"""
        self.log("Phase 2: Checking participant_access default values...")
        
        if not self.participant_token or not self.session_id:
            self.log("❌ Missing participant token or session ID", "ERROR")
            return False
            
        headers = {'Authorization': f'Bearer {self.participant_token}'}
        
        try:
            response = self.session.get(f"{BASE_URL}/participant-access/{self.session_id}", headers=headers)
            
            if response.status_code == 200:
                access = response.json()
                self.log(f"✅ Participant access retrieved successfully. ID: {access.get('id')}")
                
                # Check default values
                expected_defaults = {
                    'can_access_pre_test': False,
                    'can_access_post_test': False,
                    'can_access_checklist': False,
                    'can_access_feedback': False,
                    'pre_test_completed': False,
                    'post_test_completed': False,
                    'checklist_submitted': False,
                    'feedback_submitted': False
                }
                
                all_defaults_correct = True
                for key, expected_value in expected_defaults.items():
                    actual_value = access.get(key)
                    if actual_value != expected_value:
                        self.log(f"❌ {key}: expected {expected_value}, got {actual_value}", "ERROR")
                        all_defaults_correct = False
                    else:
                        self.log(f"✅ {key}: {actual_value} (correct default)")
                
                if all_defaults_correct:
                    self.log("✅ All participant_access default values are correct")
                    return True
                else:
                    self.log("❌ Some participant_access default values are incorrect", "ERROR")
                    return False
                    
            else:
                self.log(f"❌ Get participant access failed: {response.status_code} - {response.text}", "ERROR")
                return False
                
        except Exception as e:
            self.log(f"❌ Get participant access error: {str(e)}", "ERROR")
            return False
    
    def release_feedback(self):
        """Phase 3: As admin, release feedback for the session"""
        self.log("Phase 3: Releasing feedback as admin...")
        
        if not self.admin_token or not self.session_id:
            self.log("❌ Missing admin token or session ID", "ERROR")
            return False
            
        headers = {'Authorization': f'Bearer {self.admin_token}'}
        
        try:
            response = self.session.post(f"{BASE_URL}/sessions/{self.session_id}/release-feedback", headers=headers)
            
            if response.status_code == 200:
                data = response.json()
                self.log(f"✅ Feedback released successfully. Message: {data.get('message', 'No message')}")
                return True
            else:
                self.log(f"❌ Release feedback failed: {response.status_code} - {response.text}", "ERROR")
                return False
                
        except Exception as e:
            self.log(f"❌ Release feedback error: {str(e)}", "ERROR")
            return False
    
    def submit_feedback(self):
        """Phase 3: As participant, submit feedback"""
        self.log("Phase 3: Submitting feedback as participant...")
        
        if not self.participant_token or not self.session_id or not self.program_id:
            self.log("❌ Missing participant token, session ID, or program ID", "ERROR")
            return False
            
        headers = {'Authorization': f'Bearer {self.participant_token}'}
        
        feedback_data = {
            "session_id": self.session_id,
            "program_id": self.program_id,
            "responses": [
                {"question": "Overall Training Experience", "answer": 5},
                {"question": "Training Content Quality", "answer": 4},
                {"question": "Trainer Effectiveness", "answer": 5},
                {"question": "Venue & Facilities", "answer": 4},
                {"question": "Suggestions for Improvement", "answer": "Great training overall, very informative"},
                {"question": "Additional Comments", "answer": "Thank you for the excellent training program"}
            ]
        }
        
        try:
            response = self.session.post(f"{BASE_URL}/feedback/submit", json=feedback_data, headers=headers)
            
            if response.status_code == 200:
                data = response.json()
                self.log(f"✅ Feedback submitted successfully. ID: {data.get('id')}")
                self.log(f"   Session ID: {data.get('session_id')}")
                self.log(f"   Program ID: {data.get('program_id')}")
                self.log(f"   Responses: {len(data.get('responses', []))}")
                return True
            else:
                self.log(f"❌ Feedback submission failed: {response.status_code} - {response.text}", "ERROR")
                return False
                
        except Exception as e:
            self.log(f"❌ Feedback submission error: {str(e)}", "ERROR")
            return False
    
    def verify_feedback_submitted_flag(self):
        """Phase 3: Verify participant_access.feedback_submitted = True"""
        self.log("Phase 3: Verifying feedback_submitted flag is set to True...")
        
        if not self.participant_token or not self.session_id:
            self.log("❌ Missing participant token or session ID", "ERROR")
            return False
            
        headers = {'Authorization': f'Bearer {self.participant_token}'}
        
        try:
            response = self.session.get(f"{BASE_URL}/participant-access/{self.session_id}", headers=headers)
            
            if response.status_code == 200:
                access = response.json()
                feedback_submitted = access.get('feedback_submitted')
                
                if feedback_submitted is True:
                    self.log("✅ feedback_submitted flag is correctly set to True")
                    return True
                else:
                    self.log(f"❌ feedback_submitted flag is {feedback_submitted}, expected True", "ERROR")
                    return False
                    
            else:
                self.log(f"❌ Get participant access failed: {response.status_code} - {response.text}", "ERROR")
                return False
                
        except Exception as e:
            self.log(f"❌ Get participant access error: {str(e)}", "ERROR")
            return False
    
    def generate_certificate(self):
        """Phase 4: Generate certificate - THE CRITICAL PART"""
        self.log("Phase 4: Generating certificate - THE CRITICAL PART...")
        
        if not self.admin_token or not self.session_id or not self.participant_id:
            self.log("❌ Missing admin token, session ID, or participant ID", "ERROR")
            return False
            
        headers = {'Authorization': f'Bearer {self.admin_token}'}
        
        try:
            response = self.session.post(f"{BASE_URL}/certificates/generate/{self.session_id}/{self.participant_id}", headers=headers)
            
            if response.status_code == 200:
                data = response.json()
                self.log(f"✅ Certificate generation successful!")
                
                # Check response includes certificate_url and download_url
                certificate_url = data.get('certificate_url')
                download_url = data.get('download_url')
                
                if certificate_url:
                    self.log(f"✅ certificate_url present: {certificate_url}")
                else:
                    self.log("❌ certificate_url missing from response", "ERROR")
                    return False
                
                if download_url:
                    self.log(f"✅ download_url present: {download_url}")
                    self.download_url = download_url  # Store for next test
                else:
                    self.log("❌ download_url missing from response", "ERROR")
                    return False
                
                self.log(f"   Certificate ID: {data.get('id')}")
                self.log(f"   Participant ID: {data.get('participant_id')}")
                self.log(f"   Session ID: {data.get('session_id')}")
                self.log(f"   Program Name: {data.get('program_name')}")
                
                return True
            else:
                self.log(f"❌ Certificate generation failed: {response.status_code} - {response.text}", "ERROR")
                return False
                
        except Exception as e:
            self.log(f"❌ Certificate generation error: {str(e)}", "ERROR")
            return False
    
    def download_and_verify_certificate(self):
        """Phase 4: Download certificate file and verify it's valid"""
        self.log("Phase 4: Downloading and verifying certificate file...")
        
        if not hasattr(self, 'download_url'):
            self.log("❌ No download URL available", "ERROR")
            return False
        
        try:
            # Construct full URL if download_url is relative
            download_url = self.download_url
            if download_url.startswith('/'):
                download_url = BASE_URL.replace('/api', '') + download_url
            
            self.log(f"Downloading from: {download_url}")
            
            # Download the certificate file with admin authentication
            headers = {'Authorization': f'Bearer {self.admin_token}'}
            response = self.session.get(download_url, headers=headers)
            
            if response.status_code == 200:
                self.log("✅ Certificate file downloaded successfully")
                
                # Check file size > 0
                file_size = len(response.content)
                if file_size > 0:
                    self.log(f"✅ File size is {file_size} bytes (> 0)")
                else:
                    self.log("❌ File size is 0 bytes", "ERROR")
                    return False
                
                # Save to temporary file and verify it's a valid Word document
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                    temp_file.write(response.content)
                    temp_file_path = temp_file.name
                
                try:
                    # Try to verify it's a valid .docx file by checking the file signature
                    with open(temp_file_path, 'rb') as f:
                        # .docx files are ZIP archives, so they start with PK
                        file_header = f.read(4)
                        if file_header.startswith(b'PK'):
                            self.log("✅ File appears to be a valid ZIP-based document (.docx)")
                            
                            # Additional check: try to import docx library and open the file
                            try:
                                from docx import Document
                                doc = Document(temp_file_path)
                                self.log(f"✅ File is a valid Word document with {len(doc.paragraphs)} paragraphs")
                                
                                # Check if document has content
                                has_content = any(p.text.strip() for p in doc.paragraphs)
                                if has_content:
                                    self.log("✅ Document contains text content")
                                else:
                                    self.log("⚠️  Document appears to be empty", "WARNING")
                                
                                return True
                            except ImportError:
                                self.log("⚠️  python-docx not available, but file header looks correct", "WARNING")
                                return True
                            except Exception as docx_error:
                                self.log(f"❌ Error reading Word document: {str(docx_error)}", "ERROR")
                                return False
                        else:
                            self.log(f"❌ File does not appear to be a valid .docx file (header: {file_header})", "ERROR")
                            return False
                            
                finally:
                    # Clean up temporary file
                    try:
                        os.unlink(temp_file_path)
                    except:
                        pass
                        
            else:
                self.log(f"❌ Certificate download failed: {response.status_code} - {response.text}", "ERROR")
                return False
                
        except Exception as e:
            self.log(f"❌ Certificate download/verification error: {str(e)}", "ERROR")
            return False
    
    def check_certificate_template_exists(self):
        """Verify certificate template exists at expected location"""
        self.log("Checking if certificate template exists...")
        
        template_path = "/app/backend/static/templates/certificate_template.docx"
        
        if os.path.exists(template_path):
            file_size = os.path.getsize(template_path)
            self.log(f"✅ Certificate template exists at {template_path} ({file_size} bytes)")
            return True
        else:
            self.log(f"❌ Certificate template not found at {template_path}", "ERROR")
            return False
    
    def cleanup(self):
        """Clean up created test data"""
        self.log("Cleaning up test data...")
        
        if self.admin_token:
            headers = {'Authorization': f'Bearer {self.admin_token}'}
            
            # Delete session
            if self.session_id:
                try:
                    response = self.session.delete(f"{BASE_URL}/sessions/{self.session_id}", headers=headers)
                    if response.status_code == 200:
                        self.log(f"✅ Cleaned up session: {self.session_id}")
                    else:
                        self.log(f"⚠️  Failed to cleanup session: {response.status_code}", "WARNING")
                except Exception as e:
                    self.log(f"⚠️  Error cleaning up session: {str(e)}", "WARNING")
            
            # Delete program
            if self.program_id:
                try:
                    response = self.session.delete(f"{BASE_URL}/programs/{self.program_id}", headers=headers)
                    if response.status_code == 200:
                        self.log(f"✅ Cleaned up program: {self.program_id}")
                    else:
                        self.log(f"⚠️  Failed to cleanup program: {response.status_code}", "WARNING")
                except Exception as e:
                    self.log(f"⚠️  Error cleaning up program: {str(e)}", "WARNING")
            
            # Delete company
            if self.company_id:
                try:
                    response = self.session.delete(f"{BASE_URL}/companies/{self.company_id}", headers=headers)
                    if response.status_code == 200:
                        self.log(f"✅ Cleaned up company: {self.company_id}")
                    else:
                        self.log(f"⚠️  Failed to cleanup company: {response.status_code}", "WARNING")
                except Exception as e:
                    self.log(f"⚠️  Error cleaning up company: {str(e)}", "WARNING")
    
    def run_certificate_flow_test(self):
        """Run the complete certificate generation flow test"""
        self.log("=" * 80)
        self.log("STARTING CERTIFICATE GENERATION FLOW TEST - CRITICAL TASK")
        self.log("=" * 80)
        
        test_results = []
        
        # Test sequence as specified in the review request
        tests = [
            # Phase 1: Setup (Admin)
            ("Check Certificate Template Exists", self.check_certificate_template_exists),
            ("Admin Login", self.login_admin),
            ("Create Program", self.create_program),
            ("Create Company", self.create_company),
            ("Create Participant User", self.create_participant_user),
            ("Create Session", self.create_session),
            
            # Phase 2: Participant Access
            ("Participant Login", self.login_participant),
            ("Verify Participant Access Auto-Created", self.verify_participant_access_auto_created),
            ("Check Participant Access Default Values", self.check_participant_access_defaults),
            
            # Phase 3: Feedback Submission
            ("Release Feedback (Admin)", self.release_feedback),
            ("Submit Feedback (Participant)", self.submit_feedback),
            ("Verify Feedback Submitted Flag", self.verify_feedback_submitted_flag),
            
            # Phase 4: Certificate Generation (THE CRITICAL PART)
            ("Generate Certificate", self.generate_certificate),
            ("Download and Verify Certificate File", self.download_and_verify_certificate),
        ]
        
        for test_name, test_func in tests:
            self.log(f"\n--- Running: {test_name} ---")
            try:
                result = test_func()
                test_results.append((test_name, result))
                if not result:
                    self.log(f"❌ {test_name} FAILED", "ERROR")
                    # For critical certificate generation, show exact failure point
                    if "Certificate" in test_name:
                        self.log("🚨 CRITICAL CERTIFICATE GENERATION FAILURE DETECTED", "ERROR")
                else:
                    self.log(f"✅ {test_name} PASSED")
            except Exception as e:
                self.log(f"❌ {test_name} ERROR: {str(e)}", "ERROR")
                test_results.append((test_name, False))
        
        # Cleanup
        self.log(f"\n--- Cleanup ---")
        self.cleanup()
        
        # Summary
        self.log("\n" + "=" * 80)
        self.log("CERTIFICATE GENERATION FLOW TEST RESULTS")
        self.log("=" * 80)
        
        passed = sum(1 for _, result in test_results if result)
        total = len(test_results)
        
        for test_name, result in test_results:
            status = "✅ PASS" if result else "❌ FAIL"
            self.log(f"{status} - {test_name}")
        
        self.log(f"\nOverall: {passed}/{total} tests passed")
        
        if passed == total:
            self.log("🎉 CERTIFICATE GENERATION FLOW WORKING PERFECTLY!")
            self.log("✅ All phases completed successfully:")
            self.log("   ✅ Phase 1: Admin setup (program, company, session, participant)")
            self.log("   ✅ Phase 2: Participant access verification")
            self.log("   ✅ Phase 3: Feedback submission flow")
            self.log("   ✅ Phase 4: Certificate generation and download")
            return True
        else:
            failed_tests = [name for name, result in test_results if not result]
            self.log(f"🚨 CERTIFICATE GENERATION FLOW HAS ISSUES!")
            self.log(f"❌ Failed tests: {', '.join(failed_tests)}")
            
            # Provide specific guidance based on failure point
            if any("Certificate" in test for test in failed_tests):
                self.log("🔍 CERTIFICATE GENERATION SPECIFIC FAILURES DETECTED:")
                self.log("   - Check certificate template exists at /app/backend/static/templates/certificate_template.docx")
                self.log("   - Verify certificate generation endpoint is implemented")
                self.log("   - Check file download URLs are accessible")
                self.log("   - Ensure generated files are valid .docx documents")
            
            return False

def main():
    """Main function to run the certificate generation flow test"""
    runner = CertificateTestRunner()
    success = runner.run_certificate_flow_test()
    
    if success:
        sys.exit(0)
    else:
        sys.exit(1)

if __name__ == "__main__":
    main()