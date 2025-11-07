from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.responses import FileResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional
import uuid
from datetime import datetime, timezone, timedelta
from passlib.context import CryptContext
import jwt
import random
import shutil
import subprocess
from docx import Document
from emergentintegrations.llm.chat import LlmChat, UserMessage
import json
import asyncio

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Security
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
security = HTTPBearer()
SECRET_KEY = os.environ.get('SECRET_KEY', 'your-secret-key-change-in-production')
ALGORITHM = "HS256"

# Create the main app
app = FastAPI()
api_router = APIRouter(prefix="/api")

# Static files directory
STATIC_DIR = ROOT_DIR / "static"
STATIC_DIR.mkdir(exist_ok=True)
LOGO_DIR = STATIC_DIR / "logos"
LOGO_DIR.mkdir(exist_ok=True)
CERTIFICATE_DIR = STATIC_DIR / "certificates"
CERTIFICATE_DIR.mkdir(exist_ok=True)
CERTIFICATE_PDF_DIR = STATIC_DIR / "certificates_pdf"
CERTIFICATE_PDF_DIR.mkdir(exist_ok=True)
TEMPLATE_DIR = STATIC_DIR / "templates"
TEMPLATE_DIR.mkdir(exist_ok=True)
CHECKLIST_PHOTOS_DIR = STATIC_DIR / "checklist_photos"
CHECKLIST_PHOTOS_DIR.mkdir(exist_ok=True)

# ============ MODELS ============

class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: EmailStr
    full_name: str
    id_number: str
    role: str
    company_id: Optional[str] = None
    location: Optional[str] = None
    phone_number: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    is_active: bool = True

class UserCreate(BaseModel):
    email: EmailStr
    password: str
    full_name: str
    id_number: str
    role: str
    company_id: Optional[str] = None
    location: Optional[str] = None
    phone_number: Optional[str] = None

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class TokenResponse(BaseModel):
    access_token: str
    token_type: str
    user: User

class Company(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class CompanyCreate(BaseModel):
    name: str

class CompanyUpdate(BaseModel):
    name: str

class Program(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    description: Optional[str] = None
    pass_percentage: float = 70.0
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ProgramCreate(BaseModel):
    name: str
    description: Optional[str] = None
    pass_percentage: Optional[float] = 70.0

class ProgramUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    pass_percentage: Optional[float] = None

class Session(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    program_id: str
    company_id: str
    location: str
    start_date: str
    end_date: str
    supervisor_ids: List[str] = []
    participant_ids: List[str] = []
    trainer_assignments: List[dict] = []
    coordinator_id: Optional[str] = None
    status: str = "active"  # "active" or "inactive"
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ParticipantData(BaseModel):
    email: EmailStr
    password: str
    full_name: str
    id_number: str
    phone_number: Optional[str] = None

class SupervisorData(BaseModel):
    email: EmailStr
    password: str
    full_name: str
    id_number: str
    phone_number: Optional[str] = None

class SessionCreate(BaseModel):
    name: str
    program_id: str
    company_id: str
    location: str
    start_date: str
    end_date: str
    supervisor_ids: List[str] = []
    participant_ids: List[str] = []
    participants: List[ParticipantData] = []  # New participants to create or link
    supervisors: List[SupervisorData] = []  # New supervisors to create or link
    trainer_assignments: List[dict] = []
    coordinator_id: Optional[str] = None

class ParticipantAccess(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    participant_id: str
    session_id: str
    can_access_pre_test: bool = False
    can_access_post_test: bool = False
    can_access_checklist: bool = False
    can_access_feedback: bool = False
    pre_test_completed: bool = False
    post_test_completed: bool = False
    checklist_submitted: bool = False
    feedback_submitted: bool = False

class UpdateParticipantAccess(BaseModel):
    participant_id: str
    session_id: str
    can_access_pre_test: Optional[bool] = None
    can_access_post_test: Optional[bool] = None
    can_access_checklist: Optional[bool] = None
    can_access_feedback: Optional[bool] = None

class TestQuestion(BaseModel):
    question: str
    options: List[str]
    correct_answer: int

class Test(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    program_id: str
    test_type: str
    questions: List[TestQuestion] = []
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class TestCreate(BaseModel):
    program_id: str
    test_type: str
    questions: List[TestQuestion]

class TestResult(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    test_id: str
    participant_id: str
    session_id: str
    test_type: str
    answers: List[int] = []
    score: float = 0.0
    total_questions: int = 0
    correct_answers: int = 0
    passed: bool = False
    submitted_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    question_indices: Optional[List[int]] = None  # Store original question order for shuffled tests

class TestSubmit(BaseModel):
    test_id: str
    session_id: str
    answers: List[int]
    question_indices: Optional[List[int]] = None  # Original question indices for shuffled tests

class ChecklistTemplate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    program_id: str
    items: List[str] = []
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ChecklistTemplateCreate(BaseModel):
    program_id: str
    items: List[str]

class VehicleChecklist(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    participant_id: str
    session_id: str
    interval: str
    checklist_items: List[dict] = []
    submitted_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    verified_by: Optional[str] = None
    verified_at: Optional[datetime] = None
    verification_status: str = "pending"

class ChecklistSubmit(BaseModel):
    session_id: str
    interval: str
    checklist_items: List[dict]

class ChecklistVerify(BaseModel):
    checklist_id: str
    status: str
    comments: Optional[str] = None

class VehicleDetails(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    participant_id: str
    session_id: str
    vehicle_model: str
    registration_number: str
    roadtax_expiry: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class VehicleDetailsSubmit(BaseModel):
    session_id: str
    vehicle_model: str
    registration_number: str
    roadtax_expiry: str

class TrainingReport(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    session_id: str
    coordinator_id: str
    group_photo: Optional[str] = None
    theory_photo_1: Optional[str] = None
    theory_photo_2: Optional[str] = None
    practical_photo_1: Optional[str] = None
    practical_photo_2: Optional[str] = None
    practical_photo_3: Optional[str] = None
    additional_notes: Optional[str] = None
    status: str = "draft"  # draft, submitted
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    submitted_at: Optional[datetime] = None

class TrainingReportCreate(BaseModel):
    session_id: str
    group_photo: Optional[str] = None
    theory_photo_1: Optional[str] = None
    theory_photo_2: Optional[str] = None
    practical_photo_1: Optional[str] = None
    practical_photo_2: Optional[str] = None
    practical_photo_3: Optional[str] = None
    additional_notes: Optional[str] = None
    status: str = "draft"

class Attendance(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    participant_id: str
    session_id: str
    date: str
    clock_in: Optional[str] = None
    clock_out: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class AttendanceClockIn(BaseModel):
    session_id: str

class AttendanceClockOut(BaseModel):
    session_id: str

# Helper function to convert DOCX to PDF
def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    """Convert DOCX to PDF using LibreOffice"""
    try:
        # Use LibreOffice in headless mode to convert DOCX to PDF
        subprocess.run([
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(pdf_path.parent),
            str(docx_path)
        ], check=True, capture_output=True, timeout=30)
        return True
    except Exception as e:
        logging.error(f"PDF conversion failed: {str(e)}")
        return False

class ChecklistItem(BaseModel):
    item: str
    status: str  # "good", "needs_repair"
    comments: str = ""
    photo_url: Optional[str] = None

class TrainerChecklistSubmit(BaseModel):
    participant_id: str
    session_id: str
    items: List[ChecklistItem]

class FeedbackQuestion(BaseModel):
    question: str
    type: str  # "rating" or "text"
    required: bool = True

class FeedbackTemplate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    program_id: str
    questions: List[FeedbackQuestion]
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class FeedbackTemplateCreate(BaseModel):
    program_id: str
    questions: List[FeedbackQuestion]

class CourseFeedback(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    participant_id: str
    session_id: str
    program_id: str
    responses: List[dict]  # [{"question": str, "answer": str/int}]
    submitted_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class FeedbackSubmit(BaseModel):
    session_id: str
    program_id: str
    responses: List[dict]  # [{"question": str, "answer": str/int}]

class Certificate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    participant_id: str
    session_id: str
    program_name: str
    issue_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    certificate_url: Optional[str] = None

class Settings(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = "app_settings"
    logo_url: Optional[str] = None
    company_name: str = "Malaysian Defensive Driving and Riding Centre Sdn Bhd"
    primary_color: str = "#3b82f6"
    secondary_color: str = "#6366f1"
    footer_text: str = ""
    certificate_template_url: Optional[str] = None
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class SettingsUpdate(BaseModel):
    company_name: Optional[str] = None
    primary_color: Optional[str] = None
    secondary_color: Optional[str] = None
    footer_text: Optional[str] = None

# ============ HELPER FUNCTIONS ============

def hash_password(password: str) -> str:
    return pwd_context.hash(password)

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)

def create_access_token(data: dict, expires_delta: timedelta = timedelta(days=7)):
    to_encode = data.copy()
    expire = datetime.now(timezone.utc) + expires_delta
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        token = credentials.credentials
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = payload.get("sub")
        if user_id is None:
            raise HTTPException(status_code=401, detail="Invalid token")
        
        user_doc = await db.users.find_one({"id": user_id}, {"_id": 0})
        if not user_doc:
            raise HTTPException(status_code=401, detail="User not found")
        
        if isinstance(user_doc.get('created_at'), str):
            user_doc['created_at'] = datetime.fromisoformat(user_doc['created_at'])
        
        return User(**user_doc)
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid token")

async def get_or_create_participant_access(participant_id: str, session_id: str):
    access_doc = await db.participant_access.find_one(
        {"participant_id": participant_id, "session_id": session_id},
        {"_id": 0}
    )
    
    if not access_doc:
        access_obj = ParticipantAccess(
            participant_id=participant_id,
            session_id=session_id
        )
        doc = access_obj.model_dump()
        await db.participant_access.insert_one(doc)
        return access_obj
    
    return ParticipantAccess(**access_doc)

async def find_or_create_user(user_data: dict, role: str, company_id: str) -> dict:
    """
    Find existing user by fullname OR email OR id_number (any match)
    If found: update the user with new data
    If not found: create new user
    Returns: user dict with 'is_existing' flag and user data
    """
    full_name = user_data.get("full_name")
    email = user_data.get("email")
    id_number = user_data.get("id_number")
    phone_number = user_data.get("phone_number")
    
    # Search for existing user by fullname OR email OR id_number
    query = {"$or": []}
    
    if full_name:
        query["$or"].append({"full_name": full_name})
    if email:
        query["$or"].append({"email": email})
    if id_number:
        query["$or"].append({"id_number": id_number})
    
    # If no fields provided, skip search
    if not query["$or"]:
        query = None
    
    existing_user = None
    if query:
        existing_user = await db.users.find_one(query, {"_id": 0})
    
    if existing_user:
        # User found - update with new data
        update_data = {
            "email": email,
            "id_number": user_data.get("id_number"),
            "phone_number": phone_number,
            "company_id": company_id,
        }
        # Remove None values
        update_data = {k: v for k, v in update_data.items() if v is not None}
        
        await db.users.update_one(
            {"id": existing_user["id"]},
            {"$set": update_data}
        )
        
        # Return updated user data
        updated_user = await db.users.find_one({"id": existing_user["id"]}, {"_id": 0})
        if isinstance(updated_user.get('created_at'), str):
            updated_user['created_at'] = datetime.fromisoformat(updated_user['created_at'])
        
        return {
            "is_existing": True,
            "user": User(**updated_user)
        }
    else:
        # User not found - create new
        hashed_password = pwd_context.hash(user_data.get("password"))
        new_user = User(
            email=email,
            full_name=full_name,
            id_number=user_data.get("id_number"),
            role=role,
            company_id=company_id,
            phone_number=phone_number
        )
        
        user_doc = new_user.model_dump()
        user_doc["created_at"] = user_doc["created_at"].isoformat()
        user_doc["password"] = hashed_password
        
        await db.users.insert_one(user_doc)
        
        return {
            "is_existing": False,
            "user": new_user
        }

# Training Report Models
class TrainingReport(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    session_id: str
    program_id: str
    company_id: str
    generated_by: str  # coordinator_id
    content: str  # Markdown content
    status: str  # "draft" or "published"
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    published_at: Optional[datetime] = None
    published_to_supervisors: List[str] = []  # List of supervisor IDs

class ReportGenerateRequest(BaseModel):
    session_id: str

class ReportUpdateRequest(BaseModel):
    content: str

# ============ ROUTES ============

@api_router.get("/")
async def root():
    return {"message": "Defensive Driving Training API"}

# Auth Routes
@api_router.post("/auth/register", response_model=User)
async def register_user(user_data: UserCreate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can create users")
    
    existing = await db.users.find_one({"email": user_data.email}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="User already exists")
    
    hashed_pw = hash_password(user_data.password)
    user_obj = User(
        email=user_data.email,
        full_name=user_data.full_name,
        id_number=user_data.id_number,
        role=user_data.role,
        company_id=user_data.company_id,
        location=user_data.location
    )
    
    doc = user_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    doc['password'] = hashed_pw
    
    await db.users.insert_one(doc)
    return user_obj

@api_router.post("/auth/login", response_model=TokenResponse)
async def login(user_data: UserLogin):
    user_doc = await db.users.find_one({"email": user_data.email}, {"_id": 0})
    if not user_doc:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Check for both 'password' and 'hashed_password' field names
    password_hash = user_doc.get('password') or user_doc.get('hashed_password')
    if not password_hash:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    if not verify_password(user_data.password, password_hash):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    if not user_doc.get('is_active', True):
        raise HTTPException(status_code=401, detail="Account is inactive")
    
    token = create_access_token({"sub": user_doc['id']})
    
    if isinstance(user_doc.get('created_at'), str):
        user_doc['created_at'] = datetime.fromisoformat(user_doc['created_at'])
    
    user_doc.pop('password', None)
    user_doc.pop('hashed_password', None)
    user = User(**user_doc)
    
    return TokenResponse(access_token=token, token_type="bearer", user=user)

@api_router.get("/auth/me", response_model=User)
async def get_me(current_user: User = Depends(get_current_user)):
    return current_user

class ForgotPasswordRequest(BaseModel):
    email: EmailStr

class ResetPasswordRequest(BaseModel):
    email: EmailStr
    new_password: str

@api_router.post("/auth/forgot-password")
async def forgot_password(request: ForgotPasswordRequest):
    """
    Simple forgot password endpoint that checks if user exists
    In production, this would send an email with reset link
    """
    user_doc = await db.users.find_one({"email": request.email}, {"_id": 0})
    
    # Always return success to prevent email enumeration
    if not user_doc:
        return {"message": "If an account exists with this email, password reset instructions have been sent"}
    
    # For MVP: Return success message
    # In production: Generate token, send email with reset link
    return {"message": "If an account exists with this email, password reset instructions have been sent"}

@api_router.post("/auth/reset-password")
async def reset_password(request: ResetPasswordRequest):
    """
    Reset password for a user
    In production, this would require a valid reset token from email
    For MVP: Allow direct reset with email verification
    """
    user_doc = await db.users.find_one({"email": request.email}, {"_id": 0})
    
    if not user_doc:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Hash new password
    hashed_password = pwd_context.hash(request.new_password)
    
    # Update password
    await db.users.update_one(
        {"email": request.email},
        {"$set": {"password": hashed_password}}
    )
    
    return {"message": "Password reset successfully"}

# Company Routes
@api_router.post("/companies", response_model=Company)
async def create_company(company_data: CompanyCreate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can create companies")
    
    company_obj = Company(name=company_data.name)
    doc = company_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.companies.insert_one(doc)
    return company_obj

@api_router.get("/companies", response_model=List[Company])
async def get_companies(current_user: User = Depends(get_current_user)):
    companies = await db.companies.find({}, {"_id": 0}).to_list(1000)
    for company in companies:
        if isinstance(company.get('created_at'), str):
            company['created_at'] = datetime.fromisoformat(company['created_at'])
    return companies

@api_router.put("/companies/{company_id}", response_model=Company)
async def update_company(company_id: str, company_data: CompanyUpdate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can update companies")
    
    result = await db.companies.update_one(
        {"id": company_id},
        {"$set": company_data.model_dump()}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Company not found")
    
    company_doc = await db.companies.find_one({"id": company_id}, {"_id": 0})
    if isinstance(company_doc.get('created_at'), str):
        company_doc['created_at'] = datetime.fromisoformat(company_doc['created_at'])
    return Company(**company_doc)

@api_router.delete("/companies/{company_id}")
async def delete_company(company_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can delete companies")
    
    result = await db.companies.delete_one({"id": company_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Company not found")
    
    return {"message": "Company deleted successfully"}

# Program Routes
@api_router.post("/programs", response_model=Program)
async def create_program(program_data: ProgramCreate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can create programs")
    
    program_obj = Program(**program_data.model_dump())
    doc = program_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.programs.insert_one(doc)
    return program_obj

@api_router.get("/programs", response_model=List[Program])
async def get_programs(current_user: User = Depends(get_current_user)):
    programs = await db.programs.find({}, {"_id": 0}).to_list(1000)
    for program in programs:
        if isinstance(program.get('created_at'), str):
            program['created_at'] = datetime.fromisoformat(program['created_at'])
    return programs

@api_router.put("/programs/{program_id}", response_model=Program)
async def update_program(program_id: str, program_data: ProgramUpdate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can update programs")
    
    update_data = {k: v for k, v in program_data.model_dump().items() if v is not None}
    
    result = await db.programs.update_one(
        {"id": program_id},
        {"$set": update_data}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Program not found")
    
    program_doc = await db.programs.find_one({"id": program_id}, {"_id": 0})
    if isinstance(program_doc.get('created_at'), str):
        program_doc['created_at'] = datetime.fromisoformat(program_doc['created_at'])
    return Program(**program_doc)

@api_router.delete("/programs/{program_id}")
async def delete_program(program_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can delete programs")
    
    result = await db.programs.delete_one({"id": program_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Program not found")
    
    return {"message": "Program deleted successfully"}

# User Delete Route
@api_router.delete("/users/{user_id}")
async def delete_user(user_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can delete users")
    
    if user_id == current_user.id:
        raise HTTPException(status_code=400, detail="Cannot delete your own account")
    
    result = await db.users.delete_one({"id": user_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="User not found")
    
    return {"message": "User deleted successfully"}

# Check if user exists
@api_router.post("/users/check-exists")
async def check_user_exists(
    full_name: Optional[str] = None,
    email: Optional[str] = None,
    id_number: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    """Check if a user exists by fullname OR email OR id_number"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can check user existence")
    
    query = {"$or": []}
    
    if full_name:
        query["$or"].append({"full_name": full_name})
    if email:
        query["$or"].append({"email": email})
    if id_number:
        query["$or"].append({"id_number": id_number})
    
    if not query["$or"]:
        return {"exists": False, "user": None}
    
    existing_user = await db.users.find_one(query, {"_id": 0, "hashed_password": 0})
    
    if existing_user:
        if isinstance(existing_user.get('created_at'), str):
            existing_user['created_at'] = datetime.fromisoformat(existing_user['created_at'])
        return {
            "exists": True,
            "user": User(**existing_user)
        }
    
    return {"exists": False, "user": None}

# Session Routes
@api_router.post("/sessions")
async def create_session(session_data: SessionCreate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can create sessions")
    
    # Process new participants (find or create)
    processed_participant_ids = list(session_data.participant_ids)  # Start with existing IDs
    participant_results = []
    
    for participant_data in session_data.participants:
        result = await find_or_create_user(
            participant_data.model_dump(),
            role="participant",
            company_id=session_data.company_id
        )
        processed_participant_ids.append(result["user"].id)
        participant_results.append({
            "name": result["user"].full_name,
            "email": result["user"].email,
            "is_existing": result["is_existing"]
        })
    
    # Process new supervisors (find or create)
    processed_supervisor_ids = list(session_data.supervisor_ids)  # Start with existing IDs
    supervisor_results = []
    
    for supervisor_data in session_data.supervisors:
        result = await find_or_create_user(
            supervisor_data.model_dump(),
            role="pic_supervisor",
            company_id=session_data.company_id
        )
        processed_supervisor_ids.append(result["user"].id)
        supervisor_results.append({
            "name": result["user"].full_name,
            "email": result["user"].email,
            "is_existing": result["is_existing"]
        })
    
    # Create session with processed IDs
    session_obj = Session(
        name=session_data.name,
        program_id=session_data.program_id,
        company_id=session_data.company_id,
        location=session_data.location,
        start_date=session_data.start_date,
        end_date=session_data.end_date,
        participant_ids=processed_participant_ids,
        supervisor_ids=processed_supervisor_ids,
        trainer_assignments=session_data.trainer_assignments,
        coordinator_id=session_data.coordinator_id,
    )
    
    doc = session_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.sessions.insert_one(doc)
    
    # Create participant access records
    for participant_id in processed_participant_ids:
        await get_or_create_participant_access(participant_id, session_obj.id)
    
    return {
        "session": session_obj,
        "participant_results": participant_results,
        "supervisor_results": supervisor_results
    }

@api_router.get("/sessions", response_model=List[Session])
async def get_sessions(current_user: User = Depends(get_current_user)):
    # Non-admin users only see active sessions
    query = {}
    if current_user.role not in ["admin"]:
        query["status"] = "active"
    
    if current_user.role == "participant":
        query["participant_ids"] = current_user.id
        sessions = await db.sessions.find(query, {"_id": 0}).to_list(1000)
        
        # Auto-create participant_access records for each session
        for session in sessions:
            await get_or_create_participant_access(current_user.id, session['id'])
    elif current_user.role == "supervisor":
        query["supervisor_ids"] = current_user.id
        sessions = await db.sessions.find(query, {"_id": 0}).to_list(1000)
    else:
        sessions = await db.sessions.find(query, {"_id": 0}).to_list(1000)
    
    for session in sessions:
        if isinstance(session.get('created_at'), str):
            session['created_at'] = datetime.fromisoformat(session['created_at'])
    return sessions

@api_router.put("/sessions/{session_id}/toggle-status")
async def toggle_session_status(session_id: str, current_user: User = Depends(get_current_user)):
    """Toggle session between active and inactive (Admin only)"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can change session status")
    
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    new_status = "inactive" if session.get("status", "active") == "active" else "active"
    
    await db.sessions.update_one(
        {"id": session_id},
        {"$set": {"status": new_status}}
    )
    
    return {"message": f"Session marked as {new_status}", "status": new_status}

@api_router.get("/sessions/{session_id}", response_model=Session)
async def get_session(session_id: str, current_user: User = Depends(get_current_user)):
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    if isinstance(session.get('created_at'), str):
        session['created_at'] = datetime.fromisoformat(session['created_at'])
    
    return session

@api_router.get("/sessions/{session_id}/participants")
async def get_session_participants(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can view participants")
    
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    participants = []
    for participant_id in session['participant_ids']:
        user_doc = await db.users.find_one({"id": participant_id}, {"_id": 0, "password": 0})
        if user_doc:
            if isinstance(user_doc.get('created_at'), str):
                user_doc['created_at'] = datetime.fromisoformat(user_doc['created_at'])
            
            access = await get_or_create_participant_access(participant_id, session_id)
            
            participants.append({
                "user": user_doc,
                "access": access.model_dump()
            })
    
    return participants

@api_router.put("/sessions/{session_id}")
async def update_session(session_id: str, session_data: dict, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can update sessions")
    
    result = await db.sessions.update_one(
        {"id": session_id},
        {"$set": session_data}
    )
    
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Session not found")
    
    return {"message": "Session updated successfully"}

@api_router.delete("/sessions/{session_id}")
async def delete_session(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can delete sessions")
    
    result = await db.sessions.delete_one({"id": session_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Also delete related participant_access records
    await db.participant_access.delete_many({"session_id": session_id})
    
    return {"message": "Session deleted successfully"}

# Participant Access Routes
@api_router.post("/participant-access/update")
async def update_participant_access(access_data: UpdateParticipantAccess, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can update access")
    
    await get_or_create_participant_access(access_data.participant_id, access_data.session_id)
    
    update_fields = {}
    if access_data.can_access_pre_test is not None:
        update_fields['can_access_pre_test'] = access_data.can_access_pre_test
    if access_data.can_access_post_test is not None:
        update_fields['can_access_post_test'] = access_data.can_access_post_test
    if access_data.can_access_checklist is not None:
        update_fields['can_access_checklist'] = access_data.can_access_checklist
    if access_data.can_access_feedback is not None:
        update_fields['can_access_feedback'] = access_data.can_access_feedback
    
    await db.participant_access.update_one(
        {"participant_id": access_data.participant_id, "session_id": access_data.session_id},
        {"$set": update_fields}
    )
    
    return {"message": "Access updated successfully"}

@api_router.get("/participant-access/{session_id}")
async def get_my_access(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "participant":
        raise HTTPException(status_code=403, detail="Only participants can check access")
    
    access = await get_or_create_participant_access(current_user.id, session_id)
    return access

@api_router.get("/participant-access/session/{session_id}")
async def get_session_access(session_id: str, current_user: User = Depends(get_current_user)):
    """Get all participant access records for a session (for coordinators/admins)"""
    if current_user.role not in ["coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Access denied")
    
    access_records = await db.participant_access.find({"session_id": session_id}, {"_id": 0}).to_list(1000)
    return access_records

@api_router.post("/participant-access/session/{session_id}/toggle")
async def toggle_session_access(session_id: str, access_data: dict, current_user: User = Depends(get_current_user)):
    """Toggle access for all participants in a session (coordinator/admin)"""
    if current_user.role not in ["coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Only coordinators and admins can control access")
    
    # Get session to find all participants
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    access_type = access_data.get("access_type")
    enabled = access_data.get("enabled", False)
    
    # Map access_type to field name
    field_mapping = {
        "pre_test": "can_access_pre_test",
        "post_test": "can_access_post_test",
        "feedback": "can_access_feedback",
        "checklist": "can_access_checklist"
    }
    
    if access_type not in field_mapping:
        raise HTTPException(status_code=400, detail="Invalid access type")
    
    field_name = field_mapping[access_type]
    
    # Update all participant access records for this session
    participant_ids = session.get("participant_ids", [])
    
    for participant_id in participant_ids:
        # Ensure access record exists
        await get_or_create_participant_access(participant_id, session_id)
        
        # Update the field
        await db.participant_access.update_one(
            {"participant_id": participant_id, "session_id": session_id},
            {"$set": {field_name: enabled}}
        )
    
    status_text = "enabled" if enabled else "disabled"
    return {"message": f"{access_type} access {status_text} for {len(participant_ids)} participants"}

# Coordinator Control Routes
@api_router.post("/sessions/{session_id}/release-pre-test")
async def release_pre_test(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Only admins and coordinators can release tests")
    
    # Get session to verify it exists
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Update all participant access records for this session
    result = await db.participant_access.update_many(
        {"session_id": session_id},
        {"$set": {"can_access_pre_test": True}}
    )
    
    return {"message": f"Pre-test released to {result.modified_count} participants"}

@api_router.post("/sessions/{session_id}/release-post-test")
async def release_post_test(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Only admins and coordinators can release tests")
    
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    result = await db.participant_access.update_many(
        {"session_id": session_id},
        {"$set": {"can_access_post_test": True}}
    )
    
    return {"message": f"Post-test released to {result.modified_count} participants"}

@api_router.post("/sessions/{session_id}/release-feedback")
async def release_feedback(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Only admins and coordinators can release feedback")
    
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    result = await db.participant_access.update_many(
        {"session_id": session_id},
        {"$set": {"can_access_feedback": True}}
    )
    
    return {"message": f"Feedback form released to {result.modified_count} participants"}

@api_router.get("/sessions/{session_id}/status")
async def get_session_status(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role not in ["admin", "coordinator", "trainer"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    # Get session
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Get all participant access records
    access_records = await db.participant_access.find({"session_id": session_id}, {"_id": 0}).to_list(1000)
    
    total_participants = len(access_records)
    pre_test_released = any(a.get('can_access_pre_test', False) for a in access_records)
    post_test_released = any(a.get('can_access_post_test', False) for a in access_records)
    feedback_released = any(a.get('can_access_feedback', False) for a in access_records)
    
    pre_test_completed = sum(1 for a in access_records if a.get('pre_test_completed', False))
    post_test_completed = sum(1 for a in access_records if a.get('post_test_completed', False))
    feedback_submitted = sum(1 for a in access_records if a.get('feedback_submitted', False))
    
    return {
        "session_id": session_id,
        "session_name": session.get('name', ''),
        "total_participants": total_participants,
        "pre_test": {
            "released": pre_test_released,
            "completed": pre_test_completed
        },
        "post_test": {
            "released": post_test_released,
            "completed": post_test_completed
        },
        "feedback": {
            "released": feedback_released,
            "submitted": feedback_submitted
        }
    }

@api_router.get("/sessions/{session_id}/results-summary")
async def get_results_summary(session_id: str, current_user: User = Depends(get_current_user)):
    # Check if user has permission (admin, coordinator, or chief trainer)
    if current_user.role not in ["admin", "coordinator"]:
        # Check if trainer is chief trainer for this session
        if current_user.role == "trainer":
            session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
            if not session:
                raise HTTPException(status_code=404, detail="Session not found")
            
            # Check if user is a chief trainer in this session
            is_chief = any(
                t.get('trainer_id') == current_user.id and t.get('role') == 'chief'
                for t in session.get('trainer_assignments', [])
            )
            if not is_chief:
                raise HTTPException(status_code=403, detail="Only chief trainers can view results")
        else:
            raise HTTPException(status_code=403, detail="Unauthorized")
    
    # Get all participants in the session
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    participant_ids = session.get('participant_ids', [])
    
    # Get participant details
    participants = await db.users.find(
        {"id": {"$in": participant_ids}},
        {"_id": 0, "password": 0}
    ).to_list(1000)
    
    # Get test results for all participants
    test_results = await db.test_results.find(
        {"session_id": session_id},
        {"_id": 0}
    ).to_list(1000)
    
    # Get feedback for all participants
    feedbacks = await db.feedbacks.find(
        {"session_id": session_id},
        {"_id": 0}
    ).to_list(1000)
    
    # Build summary
    summary = []
    for participant in participants:
        p_results = [r for r in test_results if r['participant_id'] == participant['id']]
        p_feedback = next((f for f in feedbacks if f['participant_id'] == participant['id']), None)
        
        pre_test = next((r for r in p_results if r['test_type'] == 'pre'), None)
        post_test = next((r for r in p_results if r['test_type'] == 'post'), None)
        
        summary.append({
            "participant": {
                "id": participant['id'],
                "name": participant['full_name'],
                "email": participant['email']
            },
            "pre_test": {
                "completed": pre_test is not None,
                "score": pre_test['score'] if pre_test else 0,
                "correct": pre_test['correct_answers'] if pre_test else 0,
                "total": pre_test['total_questions'] if pre_test else 0,
                "passed": pre_test['passed'] if pre_test else False,
                "result_id": pre_test['id'] if pre_test else None
            },
            "post_test": {
                "completed": post_test is not None,
                "score": post_test['score'] if post_test else 0,
                "correct": post_test['correct_answers'] if post_test else 0,
                "total": post_test['total_questions'] if post_test else 0,
                "passed": post_test['passed'] if post_test else False,
                "result_id": post_test['id'] if post_test else None
            },
            "feedback_submitted": p_feedback is not None
        })
    
    return {
        "session_id": session_id,
        "session_name": session.get('name', ''),
        "program_id": session.get('program_id', ''),
        "participants": summary
    }

# User Routes
@api_router.get("/users", response_model=List[User])
async def get_users(role: Optional[str] = None, current_user: User = Depends(get_current_user)):
    if current_user.role not in ["admin", "supervisor", "coordinator"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    query = {}
    if role:
        query["role"] = role
    
    users = await db.users.find(query, {"_id": 0, "password": 0}).to_list(1000)
    for user in users:
        if isinstance(user.get('created_at'), str):
            user['created_at'] = datetime.fromisoformat(user['created_at'])
    return users

@api_router.get("/users/{user_id}", response_model=User)
async def get_user(user_id: str, current_user: User = Depends(get_current_user)):
    # Allow access if: admin, supervisor, or the user themselves
    if current_user.role not in ["admin", "supervisor", "trainer", "coordinator"] and current_user.id != user_id:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    user = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    if isinstance(user.get('created_at'), str):
        user['created_at'] = datetime.fromisoformat(user['created_at'])
    
    return user

# Test Routes
@api_router.post("/tests", response_model=Test)
async def create_test(test_data: TestCreate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can create tests")
    
    test_obj = Test(**test_data.model_dump())
    doc = test_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.tests.insert_one(doc)
    return test_obj

@api_router.get("/tests/program/{program_id}", response_model=List[Test])
async def get_tests_by_program(program_id: str, current_user: User = Depends(get_current_user)):
    tests = await db.tests.find({"program_id": program_id}, {"_id": 0}).to_list(100)
    for test in tests:
        if isinstance(test.get('created_at'), str):
            test['created_at'] = datetime.fromisoformat(test['created_at'])
    return tests

@api_router.delete("/tests/{test_id}")
async def delete_test(test_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can delete tests")
    
    result = await db.tests.delete_one({"id": test_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Test not found")
    
    return {"message": "Test deleted successfully"}

@api_router.get("/sessions/{session_id}/tests/available")
async def get_available_tests(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "participant":
        raise HTTPException(status_code=403, detail="Only participants can access this")
    
    # Get session
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Get participant access
    access = await get_or_create_participant_access(current_user.id, session_id)
    
    # Get tests for the session's program
    tests = await db.tests.find({"program_id": session['program_id']}, {"_id": 0}).to_list(10)
    
    available_tests = []
    for test in tests:
        if isinstance(test.get('created_at'), str):
            test['created_at'] = datetime.fromisoformat(test['created_at'])
        
        test_type = test['test_type']
        can_access = False
        is_completed = False
        
        if test_type == "pre":
            can_access = access.can_access_pre_test
            is_completed = access.pre_test_completed
        elif test_type == "post":
            can_access = access.can_access_post_test
            is_completed = access.post_test_completed
        
        if can_access and not is_completed:
            # Don't send correct answers to participant
            test_copy = test.copy()
            questions = test['questions'].copy()
            
            # Shuffle post-test questions
            if test_type == "post":
                random.shuffle(questions)
            
            test_copy['questions'] = [
                {
                    'question': q['question'],
                    'options': q['options']
                }
                for q in questions
            ]
            available_tests.append(test_copy)
    
    return available_tests

@api_router.get("/tests/{test_id}")
async def get_test(test_id: str, current_user: User = Depends(get_current_user)):
    test_doc = await db.tests.find_one({"id": test_id}, {"_id": 0})
    if not test_doc:
        raise HTTPException(status_code=404, detail="Test not found")
    
    if isinstance(test_doc.get('created_at'), str):
        test_doc['created_at'] = datetime.fromisoformat(test_doc['created_at'])
    
    # Make a copy of questions for shuffling
    questions = test_doc['questions'].copy()
    
    # Shuffle post-test questions for participants
    if current_user.role == "participant" and test_doc['test_type'] == "post":
        random.shuffle(questions)
    
    # Don't send correct answers to participants before submission
    if current_user.role == "participant":
        test_doc['questions'] = [
            {
                'question': q['question'],
                'options': q['options'],
                'original_index': test_doc['questions'].index(q)  # Track original position
            }
            for q in questions
        ]
    else:
        test_doc['questions'] = questions
    
    return test_doc

@api_router.post("/tests/submit", response_model=TestResult)
async def submit_test(submission: TestSubmit, current_user: User = Depends(get_current_user)):
    if current_user.role != "participant":
        raise HTTPException(status_code=403, detail="Only participants can submit tests")
    
    test_doc = await db.tests.find_one({"id": submission.test_id}, {"_id": 0})
    if not test_doc:
        raise HTTPException(status_code=404, detail="Test not found")
    
    program_doc = await db.programs.find_one({"id": test_doc['program_id']}, {"_id": 0})
    pass_percentage = program_doc.get('pass_percentage', 70.0) if program_doc else 70.0
    
    questions = test_doc['questions']
    
    # Ensure both are integers for comparison
    correct = 0
    for i, ans in enumerate(submission.answers):
        if i < len(questions):
            # If question_indices provided (shuffled test), use original index
            if submission.question_indices and i < len(submission.question_indices):
                original_idx = submission.question_indices[i]
            else:
                original_idx = i
            
            if original_idx < len(questions):
                submitted_answer = int(ans)
                correct_answer = int(questions[original_idx]['correct_answer'])
                if submitted_answer == correct_answer:
                    correct += 1
    
    score = (correct / len(questions)) * 100 if questions else 0
    passed = score >= pass_percentage
    
    result_obj = TestResult(
        test_id=submission.test_id,
        participant_id=current_user.id,
        session_id=submission.session_id,
        test_type=test_doc['test_type'],
        answers=submission.answers,
        score=score,
        total_questions=len(questions),
        correct_answers=correct,
        passed=passed,
        question_indices=submission.question_indices  # Store the shuffled order
    )
    
    doc = result_obj.model_dump()
    doc['submitted_at'] = doc['submitted_at'].isoformat()
    
    await db.test_results.insert_one(doc)
    
    update_field = 'pre_test_completed' if test_doc['test_type'] == 'pre' else 'post_test_completed'
    await db.participant_access.update_one(
        {"participant_id": current_user.id, "session_id": submission.session_id},
        {"$set": {update_field: True}}
    )
    
    return result_obj

@api_router.get("/tests/results/participant/{participant_id}", response_model=List[TestResult])
async def get_participant_results(participant_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role == "participant" and current_user.id != participant_id:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    results = await db.test_results.find({"participant_id": participant_id}, {"_id": 0}).to_list(100)
    for result in results:
        if isinstance(result.get('submitted_at'), str):
            result['submitted_at'] = datetime.fromisoformat(result['submitted_at'])
    return results

@api_router.get("/tests/results/session/{session_id}")
async def get_session_test_results(session_id: str, current_user: User = Depends(get_current_user)):
    """Get all test results for a session (for coordinators/admins)"""
    if current_user.role not in ["coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Access denied")
    
    results = await db.test_results.find({"session_id": session_id}, {"_id": 0}).to_list(1000)
    
    for result in results:
        if isinstance(result.get('submitted_at'), str):
            result['submitted_at'] = datetime.fromisoformat(result['submitted_at'])
    
    return results

@api_router.get("/tests/results/{result_id}")
async def get_test_result_detail(result_id: str, current_user: User = Depends(get_current_user)):
    result = await db.test_results.find_one({"id": result_id}, {"_id": 0})
    if not result:
        raise HTTPException(status_code=404, detail="Test result not found")
    
    # Participants can only see their own results
    if current_user.role == "participant" and result['participant_id'] != current_user.id:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    if isinstance(result.get('submitted_at'), str):
        result['submitted_at'] = datetime.fromisoformat(result['submitted_at'])
    
    # Get the test questions with correct answers
    test = await db.tests.find_one({"id": result['test_id']}, {"_id": 0})
    if test:
        questions = test['questions']
        
        # If question_indices exists (shuffled test), reorder questions to match participant's view
        if result.get('question_indices'):
            reordered_questions = []
            for idx in result['question_indices']:
                if idx < len(questions):
                    reordered_questions.append(questions[idx])
            result['test_questions'] = reordered_questions
        else:
            result['test_questions'] = questions
    
    return result

# Checklist Template Routes
@api_router.post("/checklist-templates", response_model=ChecklistTemplate)
async def create_checklist_template(template_data: ChecklistTemplateCreate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can create checklist templates")
    
    existing = await db.checklist_templates.find_one({"program_id": template_data.program_id}, {"_id": 0})
    if existing:
        await db.checklist_templates.update_one(
            {"program_id": template_data.program_id},
            {"$set": {"items": template_data.items}}
        )
        existing['items'] = template_data.items
        if isinstance(existing.get('created_at'), str):
            existing['created_at'] = datetime.fromisoformat(existing['created_at'])
        return ChecklistTemplate(**existing)
    
    template_obj = ChecklistTemplate(**template_data.model_dump())
    doc = template_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.checklist_templates.insert_one(doc)
    return template_obj

@api_router.get("/checklist-templates", response_model=List[ChecklistTemplate])
async def get_all_checklist_templates(current_user: User = Depends(get_current_user)):
    """Get all checklist templates"""
    templates = await db.checklist_templates.find({}, {"_id": 0}).to_list(length=None)
    result = []
    for template in templates:
        if isinstance(template.get('created_at'), str):
            template['created_at'] = datetime.fromisoformat(template['created_at'])
        result.append(ChecklistTemplate(**template))
    return result

@api_router.get("/checklist-templates/program/{program_id}", response_model=ChecklistTemplate)
async def get_checklist_template(program_id: str, current_user: User = Depends(get_current_user)):
    template = await db.checklist_templates.find_one({"program_id": program_id}, {"_id": 0})
    if not template:
        return ChecklistTemplate(program_id=program_id, items=[])
    
    if isinstance(template.get('created_at'), str):
        template['created_at'] = datetime.fromisoformat(template['created_at'])
    return ChecklistTemplate(**template)

@api_router.put("/checklist-templates/{template_id}", response_model=ChecklistTemplate)
async def update_checklist_template(template_id: str, template_data: ChecklistTemplateCreate, current_user: User = Depends(get_current_user)):
    """Update a checklist template"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can update checklist templates")
    
    existing = await db.checklist_templates.find_one({"id": template_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Template not found")
    
    await db.checklist_templates.update_one(
        {"id": template_id},
        {"$set": {"items": template_data.items, "program_id": template_data.program_id}}
    )
    
    existing['items'] = template_data.items
    existing['program_id'] = template_data.program_id
    if isinstance(existing.get('created_at'), str):
        existing['created_at'] = datetime.fromisoformat(existing['created_at'])
    
    return ChecklistTemplate(**existing)

@api_router.delete("/checklist-templates/{template_id}")
async def delete_checklist_template(template_id: str, current_user: User = Depends(get_current_user)):
    """Delete a checklist template"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can delete checklist templates")
    
    result = await db.checklist_templates.delete_one({"id": template_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Template not found")
    
    return {"message": "Template deleted successfully"}

# Vehicle Details Routes
@api_router.post("/vehicle-details/submit", response_model=VehicleDetails)
async def submit_vehicle_details(vehicle_data: VehicleDetailsSubmit, current_user: User = Depends(get_current_user)):
    if current_user.role != "participant":
        raise HTTPException(status_code=403, detail="Only participants can submit vehicle details")
    
    # Check if already exists
    existing = await db.vehicle_details.find_one({
        "participant_id": current_user.id,
        "session_id": vehicle_data.session_id
    }, {"_id": 0})
    
    if existing:
        # Update existing
        await db.vehicle_details.update_one(
            {"participant_id": current_user.id, "session_id": vehicle_data.session_id},
            {"$set": {
                "vehicle_model": vehicle_data.vehicle_model,
                "registration_number": vehicle_data.registration_number,
                "roadtax_expiry": vehicle_data.roadtax_expiry
            }}
        )
        existing.update(vehicle_data.model_dump())
        if isinstance(existing.get('created_at'), str):
            existing['created_at'] = datetime.fromisoformat(existing['created_at'])
        return VehicleDetails(**existing)
    
    vehicle_obj = VehicleDetails(
        participant_id=current_user.id,
        session_id=vehicle_data.session_id,
        vehicle_model=vehicle_data.vehicle_model,
        registration_number=vehicle_data.registration_number,
        roadtax_expiry=vehicle_data.roadtax_expiry
    )
    
    doc = vehicle_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.vehicle_details.insert_one(doc)
    return vehicle_obj

@api_router.get("/vehicle-details/{session_id}/{participant_id}")
async def get_vehicle_details(session_id: str, participant_id: str, current_user: User = Depends(get_current_user)):
    vehicle = await db.vehicle_details.find_one({
        "participant_id": participant_id,
        "session_id": session_id
    }, {"_id": 0})
    
    if not vehicle:
        return None
    
    if isinstance(vehicle.get('created_at'), str):
        vehicle['created_at'] = datetime.fromisoformat(vehicle['created_at'])
    return vehicle

# Attendance Routes
@api_router.post("/attendance/clock-in")
async def clock_in(attendance_data: AttendanceClockIn, current_user: User = Depends(get_current_user)):
    if current_user.role != "participant":
        raise HTTPException(status_code=403, detail="Only participants can clock in")
    
    today = datetime.now(timezone.utc).date().isoformat()
    now = datetime.now(timezone.utc).strftime("%H:%M:%S")
    
    # Check if already clocked in today
    existing = await db.attendance.find_one({
        "participant_id": current_user.id,
        "session_id": attendance_data.session_id,
        "date": today
    }, {"_id": 0})
    
    if existing and existing.get('clock_in'):
        raise HTTPException(status_code=400, detail="Already clocked in today")
    
    if existing:
        # Update existing
        await db.attendance.update_one(
            {"id": existing['id']},
            {"$set": {"clock_in": now}}
        )
        return {"message": "Clocked in successfully", "time": now}
    
    # Create new
    attendance_obj = Attendance(
        participant_id=current_user.id,
        session_id=attendance_data.session_id,
        date=today,
        clock_in=now
    )
    
    doc = attendance_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    await db.attendance.insert_one(doc)
    
    return {"message": "Clocked in successfully", "time": now}

@api_router.post("/attendance/clock-out")
async def clock_out(attendance_data: AttendanceClockOut, current_user: User = Depends(get_current_user)):
    if current_user.role != "participant":
        raise HTTPException(status_code=403, detail="Only participants can clock out")
    
    today = datetime.now(timezone.utc).date().isoformat()
    now = datetime.now(timezone.utc).strftime("%H:%M:%S")
    
    existing = await db.attendance.find_one({
        "participant_id": current_user.id,
        "session_id": attendance_data.session_id,
        "date": today
    }, {"_id": 0})
    
    if not existing or not existing.get('clock_in'):
        raise HTTPException(status_code=400, detail="Please clock in first")
    
    if existing.get('clock_out'):
        raise HTTPException(status_code=400, detail="Already clocked out today")
    
    await db.attendance.update_one(
        {"id": existing['id']},
        {"$set": {"clock_out": now}}
    )
    
    return {"message": "Clocked out successfully", "time": now}

@api_router.get("/attendance/{session_id}/{participant_id}")
async def get_attendance(session_id: str, participant_id: str, current_user: User = Depends(get_current_user)):
    attendance_records = await db.attendance.find({
        "participant_id": participant_id,
        "session_id": session_id
    }, {"_id": 0}).to_list(100)
    
    for record in attendance_records:
        if isinstance(record.get('created_at'), str):
            record['created_at'] = datetime.fromisoformat(record['created_at'])
    
    return attendance_records

@api_router.get("/attendance/session/{session_id}")
async def get_session_attendance(session_id: str, current_user: User = Depends(get_current_user)):
    """Get all attendance records for a session (for supervisors/coordinators)"""
    if current_user.role not in ["pic_supervisor", "coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Get session to verify access
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Get all attendance records for the session
    attendance_records = await db.attendance.find({"session_id": session_id}, {"_id": 0}).to_list(1000)
    
    # Get participant details
    participant_ids = list(set([r['participant_id'] for r in attendance_records]))
    participants = await db.users.find({"id": {"$in": participant_ids}}, {"_id": 0}).to_list(1000)
    participant_map = {p['id']: p for p in participants}
    
    # Enrich attendance records with participant info
    for record in attendance_records:
        if isinstance(record.get('created_at'), str):
            record['created_at'] = datetime.fromisoformat(record['created_at'])
        participant = participant_map.get(record['participant_id'])
        if participant:
            record['participant_name'] = participant['full_name']
            record['participant_email'] = participant['email']
    
    return attendance_records

# Training Report Routes
@api_router.post("/training-reports", response_model=TrainingReport)
async def create_training_report(report_data: TrainingReportCreate, current_user: User = Depends(get_current_user)):
    """Create or update training completion report (coordinator only)"""
    if current_user.role != "coordinator":
        raise HTTPException(status_code=403, detail="Only coordinators can create training reports")
    
    # Check if report already exists for this session
    existing = await db.training_reports.find_one({"session_id": report_data.session_id}, {"_id": 0})
    
    if existing:
        # Update existing report
        update_data = report_data.model_dump()
        if update_data['status'] == 'submitted':
            update_data['submitted_at'] = datetime.now(timezone.utc).isoformat()
        
        await db.training_reports.update_one(
            {"session_id": report_data.session_id},
            {"$set": update_data}
        )
        
        updated = await db.training_reports.find_one({"session_id": report_data.session_id}, {"_id": 0})
        if isinstance(updated.get('created_at'), str):
            updated['created_at'] = datetime.fromisoformat(updated['created_at'])
        if isinstance(updated.get('submitted_at'), str):
            updated['submitted_at'] = datetime.fromisoformat(updated['submitted_at'])
        return TrainingReport(**updated)
    
    # Create new report
    report_obj = TrainingReport(
        **report_data.model_dump(),
        coordinator_id=current_user.id
    )
    
    if report_data.status == 'submitted':
        report_obj.submitted_at = datetime.now(timezone.utc)
    
    doc = report_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    if doc.get('submitted_at'):
        doc['submitted_at'] = doc['submitted_at'].isoformat()
    
    await db.training_reports.insert_one(doc)
    return report_obj

@api_router.get("/training-reports/{session_id}", response_model=TrainingReport)
async def get_training_report(session_id: str, current_user: User = Depends(get_current_user)):
    """Get training report for a session"""
    report = await db.training_reports.find_one({"session_id": session_id}, {"_id": 0})
    
    if not report:
        # Return empty report structure
        return TrainingReport(session_id=session_id, coordinator_id=current_user.id, status="draft")
    
    if isinstance(report.get('created_at'), str):
        report['created_at'] = datetime.fromisoformat(report['created_at'])
    if isinstance(report.get('submitted_at'), str) and report.get('submitted_at'):
        report['submitted_at'] = datetime.fromisoformat(report['submitted_at'])
    
    return TrainingReport(**report)

@api_router.get("/training-reports/coordinator/{coordinator_id}")
async def get_coordinator_reports(coordinator_id: str, current_user: User = Depends(get_current_user)):
    """Get all training reports for a coordinator"""
    if current_user.role != "coordinator" and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Access denied")
    
    reports = await db.training_reports.find({"coordinator_id": coordinator_id}, {"_id": 0}).to_list(100)
    
    for report in reports:
        if isinstance(report.get('created_at'), str):
            report['created_at'] = datetime.fromisoformat(report['created_at'])
        if isinstance(report.get('submitted_at'), str) and report.get('submitted_at'):
            report['submitted_at'] = datetime.fromisoformat(report['submitted_at'])
    
    return reports

@api_router.post("/training-reports/{session_id}/generate-ai-report")
async def generate_ai_report(session_id: str, current_user: User = Depends(get_current_user)):
    """Generate AI training report using ChatGPT"""
    if current_user.role != "coordinator" and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only coordinators can generate reports")
    
    from emergentintegrations.llm.chat import LlmChat, UserMessage, FileContentWithMimeType
    from dotenv import load_dotenv
    load_dotenv()
    
    # Get session details
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Get program details
    program = await db.programs.find_one({"id": session['program_id']}, {"_id": 0})
    
    # Get company details
    company = await db.companies.find_one({"id": session['company_id']}, {"_id": 0})
    
    # Get participants count
    participant_count = len(session.get('participant_ids', []))
    
    # Get attendance records
    attendance_records = await db.attendance.find({"session_id": session_id}, {"_id": 0}).to_list(1000)
    total_attendance = len(set([r['participant_id'] for r in attendance_records]))
    
    # Get test results
    test_results = await db.test_results.find({"session_id": session_id}, {"_id": 0}).to_list(1000)
    passed_tests = len([r for r in test_results if r.get('passed', False)])
    
    # Get training report with photos
    training_report = await db.training_reports.find_one({"session_id": session_id}, {"_id": 0})
    
    # Build context for AI
    context = f"""
Generate a professional defensive driving training completion report in a structured format similar to official training documentation.

**SESSION INFORMATION:**
Program Name: {program.get('name', 'N/A') if program else 'N/A'}
Company: {company.get('name', 'N/A') if company else 'N/A'}
Training Location: {session.get('location', 'N/A')}
Training Period: {session.get('start_date', 'N/A')} to {session.get('end_date', 'N/A')}
Total Participants: {participant_count}
Attendance: {total_attendance} out of {participant_count} participants
Assessment Pass Rate: {passed_tests} out of {len(test_results)} passed

**DOCUMENTATION:**
- Group Photo: {'Attached' if training_report and training_report.get('group_photo') else 'Not provided'}
- Theory Session Photos: {2 if training_report and training_report.get('theory_photo_1') and training_report.get('theory_photo_2') else 0} photos attached
- Practical Session Photos: {3 if training_report and training_report.get('practical_photo_1') and training_report.get('practical_photo_2') and training_report.get('practical_photo_3') else 0} photos attached

**REQUIRED REPORT STRUCTURE:**

# TRAINING COMPLETION REPORT

## 1. EXECUTIVE SUMMARY
[Provide a 2-3 sentence overview of the training session]

## 2. TRAINING PROGRAM DETAILS
- Program Name: [name]
- Training Duration: [dates]
- Location: [location]
- Target Audience: [company employees]

## 3. TRAINING OBJECTIVES
[List 3-4 key objectives of the defensive driving program]

## 4. TRAINING DELIVERY
**Theory Sessions:**
[Describe theory topics covered - 2-3 sentences]

**Practical Sessions:**
[Describe hands-on activities and exercises - 2-3 sentences]

## 5. PARTICIPANT PERFORMANCE
- Total Enrolled: {participant_count}
- Attendance Rate: {round((total_attendance/participant_count)*100) if participant_count > 0 else 0}%
- Assessment Pass Rate: {round((passed_tests/len(test_results))*100) if len(test_results) > 0 else 0}%

## 6. KEY LEARNING OUTCOMES
[List 4-5 key skills/knowledge participants gained]
- 
- 
- 

## 7. TRAINING EFFECTIVENESS
[Evaluate based on attendance and pass rates - 2-3 sentences]

## 8. OBSERVATIONS & FEEDBACK
[Note any significant observations about participant engagement, questions asked, areas of difficulty]

## 9. RECOMMENDATIONS
[Provide 2-3 recommendations for future training sessions]

## 10. CONCLUSION
[Summarize the overall success of the training]

---
Report Prepared By: Training Coordinator
Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}

Please generate this report professionally with proper formatting, specific details based on the data provided, and maintain a formal tone suitable for official documentation.
"""
    
    try:
        # Initialize LLM Chat
        api_key = os.environ.get('EMERGENT_LLM_KEY', '')
        if not api_key:
            raise HTTPException(status_code=500, detail="EMERGENT_LLM_KEY not configured")
        
        chat = LlmChat(
            api_key=api_key,
            session_id=f"report_{session_id}",
            system_message="You are a professional training report writer specializing in defensive driving and road safety training programs."
        ).with_model("openai", "gpt-4o")
        
        user_message = UserMessage(text=context)
        
        # Generate report
        ai_response = await chat.send_message(user_message)
        
        return {
            "session_id": session_id,
            "generated_report": ai_response,
            "metadata": {
                "participant_count": participant_count,
                "attendance_rate": f"{total_attendance}/{participant_count}",
                "test_pass_rate": f"{passed_tests}/{len(test_results)}",
                "photos_included": bool(training_report)
            }
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate AI report: {str(e)}")

# Trainer Checklist Routes
@api_router.post("/trainer-checklist/submit")
async def submit_trainer_checklist(checklist_data: TrainerChecklistSubmit, current_user: User = Depends(get_current_user)):
    if current_user.role != "trainer":
        raise HTTPException(status_code=403, detail="Only trainers can submit checklists")
    
    # Create checklist
    checklist_obj = VehicleChecklist(
        participant_id=checklist_data.participant_id,
        session_id=checklist_data.session_id,
        interval="trainer_inspection",
        checklist_items=[item.model_dump() for item in checklist_data.items],
        verified_by=current_user.id,
        verified_at=datetime.now(timezone.utc),
        verification_status="completed"
    )
    
    doc = checklist_obj.model_dump()
    doc['submitted_at'] = doc['submitted_at'].isoformat()
    doc['verified_at'] = doc['verified_at'].isoformat()
    
    await db.vehicle_checklists.insert_one(doc)
    
    return {"message": "Checklist submitted successfully", "checklist_id": checklist_obj.id}

@api_router.get("/trainer-checklist/{session_id}/assigned-participants")
async def get_assigned_participants(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "trainer":
        raise HTTPException(status_code=403, detail="Only trainers can access this")
    
    # Get session
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Get all trainers in session
    trainer_assignments = session.get('trainer_assignments', [])
    trainers = [t['trainer_id'] for t in trainer_assignments]
    
    if not trainers:
        return []
    
    # Auto-assign participants to trainers
    participant_ids = session.get('participant_ids', [])
    total_participants = len(participant_ids)
    total_trainers = len(trainers)
    
    if total_trainers == 0:
        return []
    
    # Find chief trainer
    chief_trainers = [t['trainer_id'] for t in trainer_assignments if t.get('role') == 'chief']
    regular_trainers = [t['trainer_id'] for t in trainer_assignments if t.get('role') != 'chief']
    
    # Chief trainers get FEWER participants (supervisory role), regular trainers get MORE (do the work)
    if chief_trainers:
        # Allocate less to chief trainers (they supervise)
        total_chief = len(chief_trainers)
        total_regular = len(regular_trainers)
        
        # Give 30% to chiefs, 70% to regulars (chiefs supervise, regulars do hands-on work)
        if total_regular > 0:
            participants_for_chiefs = int(total_participants * 0.3)
            participants_for_regular = total_participants - participants_for_chiefs
            
            if current_user.id in chief_trainers:
                participants_per_chief = participants_for_chiefs // total_chief if total_chief > 0 else 0
                chief_index = chief_trainers.index(current_user.id)
                start_index = chief_index * participants_per_chief
                assigned_count = participants_per_chief
                # Distribute remainder evenly among chiefs
                if chief_index < (participants_for_chiefs % total_chief):
                    assigned_count += 1
            else:
                participants_per_regular = participants_for_regular // total_regular if total_regular > 0 else 0
                regular_index = regular_trainers.index(current_user.id)
                start_index = participants_for_chiefs + (regular_index * participants_per_regular)
                assigned_count = participants_per_regular
                # Distribute remainder evenly among regulars
                if regular_index < (participants_for_regular % total_regular):
                    assigned_count += 1
        else:
            # Only chief trainers
            participants_per_chief = total_participants // total_chief
            chief_index = chief_trainers.index(current_user.id)
            start_index = chief_index * participants_per_chief
            assigned_count = participants_per_chief
            if chief_index < (total_participants % total_chief):
                assigned_count += 1
    else:
        # No chief trainers, divide equally
        participants_per_trainer = total_participants // total_trainers
        remainder = total_participants % total_trainers
        current_trainer_index = trainers.index(current_user.id)
        start_index = current_trainer_index * participants_per_trainer
        assigned_count = participants_per_trainer
        if current_trainer_index < remainder:
            assigned_count += 1
    
    end_index = start_index + assigned_count
    assigned_participant_ids = participant_ids[start_index:end_index]
    
    # Get participant details
    participants = await db.users.find(
        {"id": {"$in": assigned_participant_ids}},
        {"_id": 0, "password": 0}
    ).to_list(100)
    
    # Get vehicle details for each
    for participant in participants:
        vehicle = await db.vehicle_details.find_one({
            "participant_id": participant['id'],
            "session_id": session_id
        }, {"_id": 0})
        participant['vehicle_details'] = vehicle
        
        # Get existing checklist
        checklist = await db.vehicle_checklists.find_one({
            "participant_id": participant['id'],
            "session_id": session_id,
            "verified_by": current_user.id
        }, {"_id": 0})
        participant['checklist'] = checklist
    
    return participants

# Vehicle Checklist Routes
@api_router.post("/checklists/submit", response_model=VehicleChecklist)
async def submit_checklist(checklist_data: ChecklistSubmit, current_user: User = Depends(get_current_user)):
    if current_user.role != "participant":
        raise HTTPException(status_code=403, detail="Only participants can submit checklists")
    
    checklist_obj = VehicleChecklist(
        participant_id=current_user.id,
        session_id=checklist_data.session_id,
        interval=checklist_data.interval,
        checklist_items=checklist_data.checklist_items
    )
    
    doc = checklist_obj.model_dump()
    doc['submitted_at'] = doc['submitted_at'].isoformat()
    if doc.get('verified_at'):
        doc['verified_at'] = doc['verified_at'].isoformat()
    
    await db.vehicle_checklists.insert_one(doc)
    
    await db.participant_access.update_one(
        {"participant_id": current_user.id, "session_id": checklist_data.session_id},
        {"$set": {"checklist_submitted": True}}
    )
    
    return checklist_obj

@api_router.get("/checklists/participant/{participant_id}")
async def get_participant_checklists(participant_id: str, current_user: User = Depends(get_current_user)):
    """Get all checklists for a participant (completed by trainers)"""
    # Allow participant themselves, trainers, coordinators, and admins
    if current_user.role not in ["trainer", "coordinator", "admin"] and current_user.id != participant_id:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    checklists = await db.vehicle_checklists.find({
        "participant_id": participant_id
    }, {"_id": 0}).to_list(1000)
    
    for checklist in checklists:
        if isinstance(checklist.get('submitted_at'), str):
            checklist['submitted_at'] = datetime.fromisoformat(checklist['submitted_at'])
        if checklist.get('verified_at') and isinstance(checklist['verified_at'], str):
            checklist['verified_at'] = datetime.fromisoformat(checklist['verified_at'])
    
    return checklists

@api_router.get("/vehicle-checklists/{session_id}/{participant_id}")
async def get_checklist(session_id: str, participant_id: str, current_user: User = Depends(get_current_user)):
    # Allow trainer, coordinator, admin, or the participant themselves
    if current_user.role not in ["trainer", "coordinator", "admin"] and current_user.id != participant_id:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    checklist = await db.vehicle_checklists.find_one({
        "participant_id": participant_id,
        "session_id": session_id
    }, {"_id": 0})
    
    if not checklist:
        raise HTTPException(status_code=404, detail="Checklist not found")
    
    if isinstance(checklist.get('submitted_at'), str):
        checklist['submitted_at'] = datetime.fromisoformat(checklist['submitted_at'])
    if checklist.get('verified_at') and isinstance(checklist['verified_at'], str):
        checklist['verified_at'] = datetime.fromisoformat(checklist['verified_at'])
    
    return checklist

@api_router.get("/checklists/participant/{participant_id}", response_model=List[VehicleChecklist])
async def get_participant_checklists(participant_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role == "participant" and current_user.id != participant_id:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    checklists = await db.vehicle_checklists.find({"participant_id": participant_id}, {"_id": 0}).to_list(100)
    for checklist in checklists:
        if isinstance(checklist.get('submitted_at'), str):
            checklist['submitted_at'] = datetime.fromisoformat(checklist['submitted_at'])
        if checklist.get('verified_at') and isinstance(checklist['verified_at'], str):
            checklist['verified_at'] = datetime.fromisoformat(checklist['verified_at'])
    return checklists

@api_router.get("/checklists/pending", response_model=List[VehicleChecklist])
async def get_pending_checklists(current_user: User = Depends(get_current_user)):
    if current_user.role != "supervisor" and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only supervisors can verify checklists")
    
    checklists = await db.vehicle_checklists.find({"verification_status": "pending"}, {"_id": 0}).to_list(100)
    for checklist in checklists:
        if isinstance(checklist.get('submitted_at'), str):
            checklist['submitted_at'] = datetime.fromisoformat(checklist['submitted_at'])
    return checklists

@api_router.post("/checklists/verify")
async def verify_checklist(verification: ChecklistVerify, current_user: User = Depends(get_current_user)):
    if current_user.role != "supervisor" and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only supervisors can verify checklists")
    
    result = await db.vehicle_checklists.update_one(
        {"id": verification.checklist_id},
        {
            "$set": {
                "verification_status": verification.status,
                "verified_by": current_user.id,
                "verified_at": datetime.now(timezone.utc).isoformat()
            }
        }
    )
    
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Checklist not found")
    
    return {"message": "Checklist verified successfully"}

# Course Feedback Routes
# Feedback Template Routes
@api_router.post("/feedback-templates", response_model=FeedbackTemplate)
async def create_feedback_template(template_data: FeedbackTemplateCreate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can create feedback templates")
    
    # Delete existing template for this program
    await db.feedback_templates.delete_many({"program_id": template_data.program_id})
    
    template_obj = FeedbackTemplate(
        program_id=template_data.program_id,
        questions=template_data.questions
    )
    
    doc = template_obj.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    
    await db.feedback_templates.insert_one(doc)
    
    return template_obj

@api_router.get("/feedback-templates/program/{program_id}")
async def get_feedback_template(program_id: str, current_user: User = Depends(get_current_user)):
    template = await db.feedback_templates.find_one({"program_id": program_id}, {"_id": 0})
    if not template:
        # Return default template instead of error
        return {
            "program_id": program_id,
            "questions": [
                {"question": "Overall Training Experience", "type": "rating", "required": True},
                {"question": "Training Content Quality", "type": "rating", "required": True},
                {"question": "Trainer Effectiveness", "type": "rating", "required": True},
                {"question": "Venue & Facilities", "type": "rating", "required": True},
                {"question": "Suggestions for Improvement", "type": "text", "required": False},
                {"question": "Additional Comments", "type": "text", "required": False}
            ]
        }
    
    if isinstance(template.get('created_at'), str):
        template['created_at'] = datetime.fromisoformat(template['created_at'])
    
    return template

@api_router.delete("/feedback-templates/{template_id}")
async def delete_feedback_template(template_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can delete feedback templates")
    
    result = await db.feedback_templates.delete_one({"id": template_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Feedback template not found")
    
    return {"message": "Feedback template deleted successfully"}

@api_router.post("/feedback/submit", response_model=CourseFeedback)
async def submit_feedback(feedback_data: FeedbackSubmit, current_user: User = Depends(get_current_user)):
    if current_user.role != "participant":
        raise HTTPException(status_code=403, detail="Only participants can submit feedback")
    
    feedback_obj = CourseFeedback(
        participant_id=current_user.id,
        session_id=feedback_data.session_id,
        program_id=feedback_data.program_id,
        responses=feedback_data.responses
    )
    
    doc = feedback_obj.model_dump()
    doc['submitted_at'] = doc['submitted_at'].isoformat()
    
    await db.course_feedback.insert_one(doc)
    
    # Ensure participant_access record exists and update feedback status
    await db.participant_access.update_one(
        {"participant_id": current_user.id, "session_id": feedback_data.session_id},
        {"$set": {"feedback_submitted": True}},
        upsert=True
    )
    
    return feedback_obj

@api_router.get("/feedback/session/{session_id}", response_model=List[CourseFeedback])
async def get_session_feedback(session_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role not in ["admin", "supervisor", "coordinator"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    feedback = await db.course_feedback.find({"session_id": session_id}, {"_id": 0}).to_list(100)
    for fb in feedback:
        if isinstance(fb.get('submitted_at'), str):
            fb['submitted_at'] = datetime.fromisoformat(fb['submitted_at'])
    return feedback

@api_router.get("/feedback/company/{company_id}")
async def get_company_feedback(company_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can view company feedback")
    
    sessions = await db.sessions.find({"company_id": company_id}, {"_id": 0}).to_list(1000)
    session_ids = [s['id'] for s in sessions]
    
    feedback = await db.course_feedback.find({"session_id": {"$in": session_ids}}, {"_id": 0}).to_list(1000)
    for fb in feedback:
        if isinstance(fb.get('submitted_at'), str):
            fb['submitted_at'] = datetime.fromisoformat(fb['submitted_at'])
    
    return feedback

# Certificate Routes
@api_router.get("/certificates/participant/{participant_id}", response_model=List[Certificate])
async def get_participant_certificates(participant_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role == "participant" and current_user.id != participant_id:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    certificates = await db.certificates.find({"participant_id": participant_id}, {"_id": 0}).to_list(100)
    for cert in certificates:
        if isinstance(cert.get('issue_date'), str):
            cert['issue_date'] = datetime.fromisoformat(cert['issue_date'])
    return certificates

# Settings Routes
@api_router.get("/settings", response_model=Settings)
async def get_settings():
    settings = await db.settings.find_one({"id": "app_settings"}, {"_id": 0})
    if not settings:
        default_settings = Settings()
        doc = default_settings.model_dump()
        doc['updated_at'] = doc['updated_at'].isoformat()
        await db.settings.insert_one(doc)
        return default_settings
    
    if isinstance(settings.get('updated_at'), str):
        settings['updated_at'] = datetime.fromisoformat(settings['updated_at'])
    return Settings(**settings)

@api_router.post("/settings/upload-logo")
async def upload_logo(file: UploadFile = File(...), current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can update settings")
    
    file_ext = file.filename.split(".")[-1]
    filename = f"logo.{file_ext}"
    file_path = LOGO_DIR / filename
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    logo_url = f"/api/static/logos/{filename}"
    
    await db.settings.update_one(
        {"id": "app_settings"},
        {"$set": {"logo_url": logo_url, "updated_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True
    )
    
    return {"logo_url": logo_url}

@api_router.put("/settings", response_model=Settings)
async def update_settings(settings_data: SettingsUpdate, current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can update settings")
    
    update_data = {k: v for k, v in settings_data.model_dump().items() if v is not None}
    update_data['updated_at'] = datetime.now(timezone.utc).isoformat()
    
    await db.settings.update_one(
        {"id": "app_settings"},
        {"$set": update_data},
        upsert=True
    )
    
    settings = await db.settings.find_one({"id": "app_settings"}, {"_id": 0})
    if isinstance(settings.get('updated_at'), str):
        settings['updated_at'] = datetime.fromisoformat(settings['updated_at'])
    return Settings(**settings)

# Certificate Template Upload
@api_router.post("/settings/upload-certificate-template")
async def upload_certificate_template(file: UploadFile = File(...), current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can upload templates")
    
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")
    
    filename = "certificate_template.docx"
    file_path = TEMPLATE_DIR / filename
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    template_url = f"/api/static/templates/{filename}"
    
    await db.settings.update_one(
        {"id": "app_settings"},
        {"$set": {"certificate_template_url": template_url, "updated_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True
    )
    
    return {"template_url": template_url, "message": "Certificate template uploaded successfully"}

# Generate Certificate
@api_router.post("/certificates/generate/{session_id}/{participant_id}")
async def generate_certificate(session_id: str, participant_id: str, current_user: User = Depends(get_current_user)):
    # Only admin can generate, or participant can generate their own
    if current_user.role != "admin" and current_user.id != participant_id:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    # Check if feedback is submitted (required for certificate)
    access = await db.participant_access.find_one(
        {"participant_id": participant_id, "session_id": session_id},
        {"_id": 0}
    )
    
    if not access:
        # Auto-create if doesn't exist
        access = await get_or_create_participant_access(participant_id, session_id)
    
    if not access.get('feedback_submitted', False):
        raise HTTPException(status_code=400, detail="Please submit feedback first. Go to your dashboard and click 'Submit Feedback' button.")
    
    # Get participant details
    participant = await db.users.find_one({"id": participant_id}, {"_id": 0})
    if not participant:
        raise HTTPException(status_code=404, detail="Participant not found")
    
    # Get session details
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Get program details
    program = await db.programs.find_one({"id": session['program_id']}, {"_id": 0})
    program_name = program['name'] if program else "Training Program"
    
    # Get company details
    company = await db.companies.find_one({"id": session['company_id']}, {"_id": 0})
    company_name = company['name'] if company else ""
    
    # Get settings for company name (already in template, no replacement needed)
    
    # Load template
    template_path = TEMPLATE_DIR / "certificate_template.docx"
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="Certificate template not found. Please upload a template first.")
    
    # Create document from template
    doc = Document(template_path)
    
    # Replace placeholders in paragraphs
    replacements = {
        '«PARTICIPANT_NAME»': participant['full_name'],
        '«IC_NUMBER»': participant['id_number'],
        '«COMPANY_NAME»': company_name,
        '«PROGRAMME NAME»': program_name,
        '<<PROGRAMME NAME>>': program_name,
        '«VENUE»': session['location'],
        '«DATE»': session['end_date']
    }
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    
    # Save as new DOCX document
    cert_filename = f"certificate_{participant_id}_{session_id}.docx"
    cert_path = CERTIFICATE_DIR / cert_filename
    doc.save(cert_path)
    
    # Convert to PDF
    pdf_filename = f"certificate_{participant_id}_{session_id}.pdf"
    pdf_path = CERTIFICATE_PDF_DIR / pdf_filename
    convert_docx_to_pdf(cert_path, pdf_path)
    
    # Store certificate record (using PDF URL)
    cert_url = f"/api/static/certificates_pdf/{pdf_filename}"
    
    # Check if certificate already exists
    existing_cert = await db.certificates.find_one({
        "participant_id": participant_id,
        "session_id": session_id
    }, {"_id": 0})
    
    if existing_cert:
        # Update existing
        await db.certificates.update_one(
            {"id": existing_cert['id']},
            {"$set": {
                "certificate_url": cert_url,
                "issue_date": datetime.now(timezone.utc).isoformat()
            }}
        )
        cert_id = existing_cert['id']
    else:
        # Create new
        cert_obj = Certificate(
            participant_id=participant_id,
            session_id=session_id,
            program_name=program_name,
            certificate_url=cert_url
        )
        doc_cert = cert_obj.model_dump()
        doc_cert['issue_date'] = doc_cert['issue_date'].isoformat()
        await db.certificates.insert_one(doc_cert)
        cert_id = cert_obj.id
    
    return {
        "certificate_id": cert_id,
        "certificate_url": cert_url,
        "download_url": f"/api/certificates/download/{cert_id}",
        "message": "Certificate generated successfully"
    }

@api_router.get("/certificates/download/{certificate_id}")
async def download_certificate(certificate_id: str, current_user: User = Depends(get_current_user)):
    cert = await db.certificates.find_one({"id": certificate_id}, {"_id": 0})
    if not cert:
        raise HTTPException(status_code=404, detail="Certificate not found")
    
    # Only participant or admin can download
    if current_user.role != "admin" and current_user.id != cert['participant_id']:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    cert_url = cert['certificate_url']
    filename = cert_url.split('/')[-1]
    
    # Check if it's a PDF or DOCX
    if filename.endswith('.pdf'):
        file_path = CERTIFICATE_PDF_DIR / filename
        media_type = 'application/pdf'
    else:
        file_path = CERTIFICATE_DIR / filename
        media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Certificate file not found")
    
    return FileResponse(
        file_path,
        media_type=media_type,
        filename=filename,
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@api_router.get("/certificates/preview/{certificate_id}")
async def preview_certificate(certificate_id: str, current_user: User = Depends(get_current_user)):
    cert = await db.certificates.find_one({"id": certificate_id}, {"_id": 0})
    if not cert:
        raise HTTPException(status_code=404, detail="Certificate not found")
    
    # Only participant or admin can preview
    if current_user.role != "admin" and current_user.id != cert['participant_id']:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    cert_url = cert['certificate_url']
    filename = cert_url.split('/')[-1]
    
    # Check if it's a PDF or DOCX
    if filename.endswith('.pdf'):
        file_path = CERTIFICATE_PDF_DIR / filename
        media_type = 'application/pdf'
    else:
        file_path = CERTIFICATE_DIR / filename
        media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Certificate file not found")
    
    # Return PDF with inline disposition for browser preview
    return FileResponse(
        file_path,
        media_type=media_type,
        headers={"Content-Disposition": f"inline; filename={filename}"}
    )

# Static files
@api_router.get("/static/logos/{filename}")
async def get_logo(filename: str):
    file_path = LOGO_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Logo not found")
    return FileResponse(file_path)

@api_router.get("/static/certificates/{filename}")
async def get_certificate(filename: str):
    file_path = CERTIFICATE_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Certificate not found")
    return FileResponse(file_path)

@api_router.get("/static/certificates_pdf/{filename}")
async def get_certificate_pdf(filename: str):
    file_path = CERTIFICATE_PDF_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Certificate PDF not found")
    return FileResponse(
        file_path, 
        media_type='application/pdf',
        headers={
            "Content-Disposition": f"attachment; filename={filename}",
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "GET",
            "X-Content-Type-Options": "nosniff"
        }
    )

@api_router.get("/static/templates/{filename}")
async def get_template(filename: str):
    file_path = TEMPLATE_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Template not found")
    return FileResponse(file_path)

@api_router.post("/checklist-photos/upload")
async def upload_checklist_photo(file: UploadFile = File(...), current_user: User = Depends(get_current_user)):
    if current_user.role != "trainer":
        raise HTTPException(status_code=403, detail="Only trainers can upload checklist photos")
    
    if not file.content_type.startswith('image/'):
        raise HTTPException(status_code=400, detail="Only image files are allowed")
    
    # Generate unique filename
    file_extension = file.filename.split('.')[-1]
    filename = f"{str(uuid.uuid4())}.{file_extension}"
    file_path = CHECKLIST_PHOTOS_DIR / filename
    
    # Save file
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    photo_url = f"/api/static/checklist-photos/{filename}"
    return {"photo_url": photo_url}

@api_router.get("/static/checklist-photos/{filename}")
async def get_checklist_photo(filename: str):
    file_path = CHECKLIST_PHOTOS_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Photo not found")
    return FileResponse(file_path)

# ============ AI REPORT GENERATION ============

async def generate_training_report_content(session_id: str, program_id: str, company_id: str) -> str:
    """Generate comprehensive training report using GPT-5"""
    
    # Gather all data
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    program = await db.programs.find_one({"id": program_id}, {"_id": 0})
    company = await db.companies.find_one({"id": company_id}, {"_id": 0})
    
    # Get all participants
    participant_ids = session.get('participant_ids', [])
    participants = []
    for pid in participant_ids:
        user = await db.users.find_one({"id": pid}, {"_id": 0})
        if user:
            participants.append(user)
    
    # Get pre-test results
    pre_tests = await db.test_results.find({
        "session_id": session_id,
        "test_type": "pre"
    }, {"_id": 0}).to_list(100)
    
    # Get post-test results
    post_tests = await db.test_results.find({
        "session_id": session_id,
        "test_type": "post"
    }, {"_id": 0}).to_list(100)
    
    # Get checklists
    checklists = await db.vehicle_checklists.find({
        "session_id": session_id
    }, {"_id": 0}).to_list(100)
    
    # Get feedback
    feedbacks = await db.course_feedback.find({
        "session_id": session_id
    }, {"_id": 0}).to_list(100)
    
    # Get attendance
    attendance = await db.attendance_records.find({
        "session_id": session_id
    }, {"_id": 0}).to_list(100)
    
    # Create participant ID to name mapping
    participant_map = {p.get('id'): p.get('full_name') for p in participants}
    
    # Build comprehensive data structure
    training_data = {
        "session": {
            "name": session.get('name'),
            "location": session.get('location'),
            "start_date": str(session.get('start_date')),
            "end_date": str(session.get('end_date'))
        },
        "program": {
            "name": program.get('name'),
            "description": program.get('description', '')
        },
        "company": {
            "name": company.get('name')
        },
        "participants": {
            "total": len(participants),
            "names": [p.get('full_name') for p in participants],
            "id_map": participant_map
        },
        "pre_test_results": {
            "total_participants": len(pre_tests),
            "average_score": sum([t.get('score', 0) for t in pre_tests]) / len(pre_tests) if pre_tests else 0,
            "pass_rate": sum([1 for t in pre_tests if t.get('passed', False)]) / len(pre_tests) * 100 if pre_tests else 0,
            "details": [{"participant": t.get('participant_id'), "score": t.get('score'), "passed": t.get('passed')} for t in pre_tests]
        },
        "post_test_results": {
            "total_participants": len(post_tests),
            "average_score": sum([t.get('score', 0) for t in post_tests]) / len(post_tests) if post_tests else 0,
            "pass_rate": sum([1 for t in post_tests if t.get('passed', False)]) / len(post_tests) * 100 if post_tests else 0,
            "improvement": (sum([t.get('score', 0) for t in post_tests]) / len(post_tests) if post_tests else 0) - (sum([t.get('score', 0) for t in pre_tests]) / len(pre_tests) if pre_tests else 0),
            "details": [{"participant": t.get('participant_id'), "score": t.get('score'), "passed": t.get('passed')} for t in post_tests]
        },
        "checklist_summary": {
            "total_checklists": len(checklists),
            "items_needing_repair": sum([len([item for item in c.get('checklist_items', []) if item.get('status') == 'needs_repair']) for c in checklists]),
            "common_issues": [],
            "details": [{"participant": c.get('participant_id'), "items": c.get('checklist_items', [])} for c in checklists]
        },
        "feedback_summary": {
            "total_responses": len(feedbacks),
            "average_ratings": {},
            "comments": [f.get('responses', {}) for f in feedbacks]
        },
        "attendance": {
            "total_records": len(attendance),
            "attendance_rate": len([a for a in attendance if a.get('clock_out_time')]) / len(attendance) * 100 if attendance else 100
        }
    }
    
    # Create prompt for GPT-5
    prompt = f"""Generate a comprehensive Defensive Driving/Riding Training Report based on the following data:

TRAINING DETAILS:
- Program: {training_data['program']['name']}
- Company: {training_data['company']['name']}
- Session: {training_data['session']['name']}
- Location: {training_data['session']['location']}
- Dates: {training_data['session']['start_date']} to {training_data['session']['end_date']}
- Total Participants: {training_data['participants']['total']}

PRE-TEST RESULTS:
- Participants Tested: {training_data['pre_test_results']['total_participants']}
- Average Score: {training_data['pre_test_results']['average_score']:.1f}%
- Pass Rate: {training_data['pre_test_results']['pass_rate']:.1f}%

POST-TEST RESULTS:
- Participants Tested: {training_data['post_test_results']['total_participants']}
- Average Score: {training_data['post_test_results']['average_score']:.1f}%
- Pass Rate: {training_data['post_test_results']['pass_rate']:.1f}%
- Improvement: {training_data['post_test_results']['improvement']:.1f}%

VEHICLE CHECKLIST FINDINGS:
- Total Checklists Completed: {training_data['checklist_summary']['total_checklists']}
- Items Needing Repair: {training_data['checklist_summary']['items_needing_repair']}

DETAILED CHECKLIST ISSUES (items marked as 'needs_repair'):
{chr(10).join([
    f"- {training_data['participants']['id_map'].get(detail['participant'], 'Unknown participant')}: " + 
    ", ".join([
        f"{item.get('item', 'Unknown item')} - {item.get('comments', 'No comment')}" 
        for item in detail['items'] 
        if item.get('status') == 'needs_repair'
    ])
    for detail in training_data['checklist_summary']['details']
    if any(item.get('status') == 'needs_repair' for item in detail['items'])
]) if training_data['checklist_summary']['items_needing_repair'] > 0 else '- No items needing repair'}

FEEDBACK:
- Total Responses: {training_data['feedback_summary']['total_responses']}

ATTENDANCE:
- Attendance Rate: {training_data['attendance']['attendance_rate']:.1f}%

Generate a professional training report with the following sections:
1. Executive Summary (2-3 paragraphs)
2. Training Overview (objectives, dates, location, participants)
3. Pre-Training Assessment (detailed analysis of pre-test results)
4. Post-Training Assessment (detailed analysis of post-test results, comparison with pre-test)
5. Vehicle Inspection Findings (summary of checklist results)
6. Participant Feedback (summary of feedback responses)
7. Key Observations and Recommendations
8. Conclusion

Use professional language, include data-driven insights, and provide actionable recommendations for the company.
Format using Markdown with proper headings and bullet points."""

    # Call GPT-5
    try:
        api_key = os.getenv('EMERGENT_LLM_KEY')
        llm = LlmChat(api_key=api_key)
        
        messages = [UserMessage(content=prompt)]
        response = llm.chat(messages=messages, model="gpt-4o")
        
        return response.content
    except Exception as e:
        logging.error(f"GPT-5 report generation failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")

@api_router.post("/reports/generate")
async def generate_report(request: ReportGenerateRequest, current_user: User = Depends(get_current_user)):
    """Generate AI training report (Coordinator only)"""
    if current_user.role not in ["coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Only coordinators can generate reports")
    
    # Get session details
    session = await db.sessions.find_one({"id": request.session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Generate report content
    content = await generate_training_report_content(
        request.session_id,
        session['program_id'],
        session['company_id']
    )
    
    # Save as draft
    report = TrainingReport(
        session_id=request.session_id,
        program_id=session['program_id'],
        company_id=session['company_id'],
        generated_by=current_user.id,
        content=content,
        status="draft"
    )
    
    await db.training_reports.insert_one(report.model_dump())
    
    return report

@api_router.get("/reports/session/{session_id}")
async def get_session_report(session_id: str, current_user: User = Depends(get_current_user)):
    """Get report for session"""
    if current_user.role not in ["coordinator", "admin", "pic_supervisor"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    report = await db.training_reports.find_one({"session_id": session_id}, {"_id": 0})
    
    # If pic_supervisor, only return published reports
    if current_user.role == "pic_supervisor":
        if not report or report.get('status') != "published":
            raise HTTPException(status_code=404, detail="No published report found")
        if current_user.id not in report.get('published_to_supervisors', []):
            raise HTTPException(status_code=403, detail="Report not published to you")
    
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    return report

@api_router.put("/reports/{report_id}")
async def update_report(report_id: str, request: ReportUpdateRequest, current_user: User = Depends(get_current_user)):
    """Update report content (draft only)"""
    if current_user.role not in ["coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Only coordinators can edit reports")
    
    report = await db.training_reports.find_one({"id": report_id}, {"_id": 0})
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    if report['status'] == "published":
        raise HTTPException(status_code=400, detail="Cannot edit published report")
    
    await db.training_reports.update_one(
        {"id": report_id},
        {"$set": {"content": request.content}}
    )
    
    return {"message": "Report updated successfully"}

@api_router.post("/reports/{report_id}/publish")
async def publish_report(report_id: str, current_user: User = Depends(get_current_user)):
    """Publish report to supervisors"""
    if current_user.role not in ["coordinator", "admin"]:
        raise HTTPException(status_code=403, detail="Only coordinators can publish reports")
    
    report = await db.training_reports.find_one({"id": report_id}, {"_id": 0})
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    # Get session to find supervisors
    session = await db.sessions.find_one({"id": report['session_id']}, {"_id": 0})
    supervisor_ids = session.get('supervisor_ids', [])
    
    await db.training_reports.update_one(
        {"id": report_id},
        {"$set": {
            "status": "published",
            "published_at": datetime.now(timezone.utc),
            "published_to_supervisors": supervisor_ids
        }}
    )
    
    return {"message": "Report published successfully", "published_to": supervisor_ids}

# ============ SUPERVISOR ENDPOINTS ============

@api_router.get("/supervisor/sessions")
async def get_supervisor_sessions(current_user: User = Depends(get_current_user)):
    """Get sessions for supervisor"""
    if current_user.role != "pic_supervisor":
        raise HTTPException(status_code=403, detail="Only supervisors can access this")
    
    # Find sessions where user is listed as supervisor
    sessions = await db.sessions.find({
        "supervisor_ids": current_user.id
    }, {"_id": 0}).to_list(100)
    
    return sessions

@api_router.get("/supervisor/attendance/{session_id}")
async def get_session_attendance(session_id: str, current_user: User = Depends(get_current_user)):
    """Get attendance for session (Supervisor)"""
    if current_user.role != "pic_supervisor":
        raise HTTPException(status_code=403, detail="Only supervisors can access this")
    
    # Verify supervisor has access to this session
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session or current_user.id not in session.get('supervisor_ids', []):
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    # Get attendance records
    attendance = await db.attendance_records.find({
        "session_id": session_id
    }, {"_id": 0}).to_list(100)
    
    # Get participant details
    for record in attendance:
        participant = await db.users.find_one({"id": record['participant_id']}, {"_id": 0, "password": 0})
        if participant:
            record['participant_name'] = participant.get('full_name')
            record['participant_email'] = participant.get('email')
    
    return attendance

# Include router
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
